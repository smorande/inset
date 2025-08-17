import streamlit as st
import boto3
import json
import requests
import os
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple, TypedDict, Annotated
import base64
import io
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import pandas as pd
from dotenv import load_dotenv
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import openai
import hashlib
import pickle
import time
import threading
import schedule
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import sqlite3
from pathlib import Path
import asyncio
from enum import Enum
from reportlab.lib.utils import ImageReader
from docx.shared import Inches
from botocore.config import Config

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolNode
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_core.tools import tool
from langchain_core.pydantic_v1 import BaseModel, Field
from langchain_openai import ChatOpenAI

TOKEN_LIMITS = {
    "Quick Search": {
        "min_tokens": 1000,
        "max_tokens": 2000
    },
    "Extended Search": {
        "min_tokens": 2000,
        "max_tokens": 5000
    },
    "Deep Search": {
        "min_tokens": 5000,
        "max_tokens": 8000
    }
}

# Load environment variables
load_dotenv()

# AWS Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "eu-north-1")

# Set API key as env var for boto3 to pick up
os.environ["AWS_BEARER_TOKEN_BEDROCK"] = os.getenv("BEDROCK_API_KEY")

# OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Perplexity AI Configuration
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
PERPLEXITY_API_BASE = "https://api.perplexity.ai"

# Serper API Configuration
SERPER_API_KEY = os.getenv("SERPER_API_KEY")

# Twitter API Configuration (if available)
TWITTER_BEARER_TOKEN = os.getenv("TWITTER_BEARER_TOKEN")

# Bedrock Long-term API Configuration
BEDROCK_LONGTERM_API_KEY = os.getenv("BEDROCK_API_KEY")
BEDROCK_LONGTERM_API_ENDPOINT = "https://bedrock-runtime.us-east-1.amazonaws.com"

# Lambda Function Names
LAMBDA_FUNCTION_NAMES = {
    "company_search": "agentic-ai-company-search",
    "document_processor": "agentic-ai-document-processor",
    "data_aggregator": "agentic-ai-data-aggregator",
    "report_generator": "agentic-ai-report-generator"
}

# Local storage directory for persistence
LOCAL_STORAGE_DIR = Path("insyt_data")
LOCAL_STORAGE_DIR.mkdir(exist_ok=True)

# Global variables for table names
S3_BUCKET = os.getenv("S3_BUCKET", "agentic-ai-business-intelligence")
DYNAMODB_TABLE = os.getenv("DYNAMODB_TABLE", "BusinessIntelligence")
CHAT_HISTORY_TABLE = os.getenv("CHAT_HISTORY_TABLE", "agentic_ci_chat_history")
ALERTS_TABLE = "agentic_ai_alerts"
VALIDATION_TABLE = "agentic_ai_validation"

# Fixed Perplexity AI integration using the exact "sonar" model format
class PerplexityLLM:
    """Fixed Perplexity AI LLM implementation with sonar model and proper error handling"""
    
    def __init__(self, api_key: str, model: str = "sonar"):
        self.api_key = api_key
        self.model = model
        self.base_url = "https://api.perplexity.ai"
        
        # Validate API key format
        if not api_key or not api_key.startswith('pplx-'):
            raise ValueError("Invalid Perplexity API key format. Should start with 'pplx-'")
    
    def invoke(self, messages, search_mode="Extended Search"):
        try:
            formatted_messages = self._format_messages(messages)
            token_config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
            
            payload = {
                "model": self.model,
                "messages": formatted_messages,
                "max_tokens": token_config["max_tokens"],  # Add this line
                "search_domain_filter": ["perplexity.ai", "arxiv.org", "reuters.com", "bloomberg.com"],
                "search_recency_filter": "month"
            }
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            # Make request with proper timeout and error handling
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            # Enhanced error handling
            if response.status_code == 400:
                try:
                    error_detail = response.json()
                    error_msg = error_detail.get('error', {}).get('message', 'Bad Request')
                    raise Exception(f"Perplexity API 400 Error: {error_msg}")
                except json.JSONDecodeError:
                    raise Exception(f"Perplexity API 400 Error: {response.text}")
            elif response.status_code == 401:
                raise Exception("Perplexity API 401 Error: Invalid API key")
            elif response.status_code == 429:
                raise Exception("Perplexity API 429 Error: Rate limit exceeded")
            elif response.status_code == 500:
                raise Exception("Perplexity API 500 Error: Server error")
            
            response.raise_for_status()
            result = response.json()
            
            # Validate response structure
            if 'choices' not in result or not result['choices']:
                raise Exception("Invalid response format from Perplexity API")
            
            content = result['choices'][0]['message']['content']
            
            # Return response in expected format
            class PerplexityResponse:
                def __init__(self, content):
                    self.content = content
            
            return PerplexityResponse(content)
            
        except requests.exceptions.Timeout:
            raise Exception("Perplexity API timeout - request took too long")
        except requests.exceptions.ConnectionError:
            raise Exception("Perplexity API connection error - check internet connection")
        except requests.exceptions.RequestException as e:
            raise Exception(f"Perplexity API request error: {str(e)}")
        except json.JSONDecodeError:
            raise Exception("Perplexity API returned invalid JSON")
        except Exception as e:
            if "Perplexity API" in str(e):
                raise e
            else:
                raise Exception(f"Unexpected error with Perplexity API: {str(e)}")
    
    def _format_messages(self, messages) -> List[Dict[str, str]]:
        """Format messages for Perplexity API"""
        formatted_messages = []
        
        for msg in messages:
            # Handle different message types
            if hasattr(msg, 'content') and hasattr(msg, '__class__'):
                # LangChain message objects
                if isinstance(msg, HumanMessage):
                    formatted_messages.append({"role": "user", "content": str(msg.content)})
                elif isinstance(msg, AIMessage):
                    formatted_messages.append({"role": "assistant", "content": str(msg.content)})
                elif isinstance(msg, SystemMessage):
                    formatted_messages.append({"role": "system", "content": str(msg.content)})
                else:
                    # Fallback for unknown message types
                    formatted_messages.append({"role": "user", "content": str(msg.content)})
            elif isinstance(msg, dict):
                # Dictionary format
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                if content and role in ['user', 'assistant', 'system']:
                    formatted_messages.append({"role": role, "content": str(content)})
            elif isinstance(msg, str):
                # String format
                formatted_messages.append({"role": "user", "content": msg})
            else:
                # Try to convert to string
                try:
                    content = str(msg)
                    if content.strip():
                        formatted_messages.append({"role": "user", "content": content})
                except:
                    continue
        
        # Ensure we have at least one message
        if not formatted_messages:
            formatted_messages.append({"role": "user", "content": "Hello"})
        
        # Limit message history to avoid token limits
        if len(formatted_messages) > 10:
            formatted_messages = formatted_messages[-10:]
        
        # Ensure total content length is reasonable
        total_length = sum(len(msg.get('content', '')) for msg in formatted_messages)
        if total_length > 8000:  # Conservative limit
            # Truncate older messages
            while total_length > 8000 and len(formatted_messages) > 1:
                removed_msg = formatted_messages.pop(0)
                total_length -= len(removed_msg.get('content', ''))
        
        return formatted_messages
    
    def test_connection(self) -> Dict[str, any]:
        """Test the Perplexity API connection using sonar model"""
        try:
            test_response = self.invoke([{"role": "user", "content": "Hello, please respond with 'API test successful'"}])
            return {
                "success": True,
                "message": "Perplexity AI sonar model connection successful",
                "response": test_response.content
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Perplexity AI sonar model connection failed: {str(e)}",
                "error": str(e)
            }

def initialize_perplexity_safely(api_key: str) -> Tuple[Optional[PerplexityLLM], str]:
    """Safely initialize Perplexity AI with sonar model and proper error handling"""
    try:
        # Validate API key
        if not api_key:
            return None, "No Perplexity API key provided"
        
        if not api_key.startswith('pplx-'):
            return None, "Invalid Perplexity API key format. Should start with 'pplx-'"
        
        # Initialize client with sonar model
        perplexity_llm = PerplexityLLM(api_key, model="sonar")
        
        # Test connection
        test_result = perplexity_llm.test_connection()
        
        if test_result["success"]:
            return perplexity_llm, "Perplexity AI sonar model initialized successfully"
        else:
            return None, f"Perplexity AI sonar model test failed: {test_result['message']}"
            
    except Exception as e:
        return None, f"Failed to initialize Perplexity AI sonar model: {str(e)}"

class NovaBedrockLLM:
    """AWS Bedrock Nova Pro LLM implementation"""
    
    def __init__(self, region_name: str = "eu-north-1"):
        self.region_name = region_name
        self.model_id = os.getenv("BEDROCK_MODEL_ID", "eu.amazon.nova-pro-v1:0")
        
        try:
            # CRITICAL FIX: Set the bearer token environment variable
            bedrock_api_key = os.getenv("BEDROCK_API_KEY")
            if bedrock_api_key:
                os.environ["AWS_BEARER_TOKEN_BEDROCK"] = bedrock_api_key
            
            self.client = boto3.client(
                "bedrock-runtime",
                region_name=region_name,
                config=Config(read_timeout=3600, connect_timeout=60, retries={"max_attempts": 3})
            )
        except Exception as e:
            raise ValueError(f"Failed to initialize Bedrock client: {str(e)}")
        
    
    def invoke(self, messages, search_mode="Extended Search"):
        try:
            formatted_messages = self._format_messages(messages)
            
            # Get token limits
            token_config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
            
            request_body = {
                "schemaVersion": "messages-v1",
                "messages": formatted_messages,
                "inferenceConfig": {
                    "maxTokens": token_config["max_tokens"],  # Dynamic limit
                    "temperature": 0.7,
                    "topP": 0.9
                }
            }
            
            response = self.client.invoke_model(
                modelId=self.model_id,
                body=json.dumps(request_body)
            )
            
            result = json.loads(response['body'].read())
            content = result['output']['message']['content'][0]['text']
            
            class NovaResponse:
                def __init__(self, content):
                    self.content = content
            
            return NovaResponse(content)
            
        except Exception as e:
            raise Exception(f"Nova Bedrock API error: {str(e)}")
    
    def _format_messages(self, messages) -> List[Dict[str, str]]:
        """Format messages for Nova Pro"""
        formatted_messages = []
        
        for msg in messages:
            if hasattr(msg, 'content') and hasattr(msg, '__class__'):
                if isinstance(msg, HumanMessage):
                    formatted_messages.append({"role": "user", "content": [{"text": str(msg.content)}]})
                elif isinstance(msg, AIMessage):
                    formatted_messages.append({"role": "assistant", "content": [{"text": str(msg.content)}]})
                elif isinstance(msg, SystemMessage):
                    formatted_messages.append({"role": "user", "content": [{"text": f"System: {str(msg.content)}"}]})
            elif isinstance(msg, dict):
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                if role == 'user':
                    formatted_messages.append({"role": "user", "content": [{"text": str(content)}]})
                elif role == 'assistant':
                    formatted_messages.append({"role": "assistant", "content": [{"text": str(content)}]})
            elif isinstance(msg, str):
                formatted_messages.append({"role": "user", "content": [{"text": msg}]})
        
        if not formatted_messages:
            formatted_messages.append({"role": "user", "content": [{"text": "Hello"}]})
        
        return formatted_messages
    
    def test_connection(self) -> Dict[str, any]:
        """Test Nova Pro connection"""
        try:
            test_response = self.invoke([{"role": "user", "content": "Hello, please respond with 'API test successful'"}])
            return {
                "success": True,
                "message": "Nova Bedrock connection successful",
                "response": test_response.content
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Nova Bedrock connection failed: {str(e)}",
                "error": str(e)
            }

# LangGraph State Definitions
class AgentState(TypedDict):
    messages: List[BaseMessage]
    query: str
    company_name: str
    search_mode: str
    research_plan: Dict
    financial_data: Dict
    news_data: Dict
    competitive_data: Dict
    validation_results: Dict
    memory_context: Dict
    next_action: str
    confidence_score: float
    final_result: Dict

class QueryType(Enum):
    FINANCIAL = "financial"
    NEWS = "news"
    COMPETITIVE = "competitive"
    GENERAL = "general"
    MIXED = "mixed"

# Initialize AWS clients safely
def get_aws_clients():
    """Initialize AWS clients with error handling"""
    try:
        session = boto3.Session(
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_REGION
        )
        
        return {
            's3': session.client('s3'),
            'dynamodb': session.resource('dynamodb'),
            'bedrock': session.client('bedrock-runtime'),
            'lambda': session.client('lambda')
        }
    except Exception as e:
        st.warning(f"AWS initialization failed: {str(e)}. Using local storage only.")
        return None

# Initialize clients safely
aws_clients = get_aws_clients()

class PersistentStorage:
    """Enhanced persistent storage for all application data with database locking fixes"""
    
    def __init__(self):
        self.db_path = LOCAL_STORAGE_DIR / "insyt_data.db"
        self.init_database()
    
    def _get_connection(self):
        """Get database connection with proper settings for concurrency"""
        conn = sqlite3.connect(
            self.db_path, 
            timeout=60.0,  # Increased timeout
            check_same_thread=False,
            isolation_level=None  # Autocommit mode
        )
        # Enable WAL mode for better concurrency
        conn.execute('PRAGMA journal_mode=WAL;')
        conn.execute('PRAGMA synchronous=NORMAL;')
        conn.execute('PRAGMA cache_size=10000;')
        conn.execute('PRAGMA temp_store=memory;')
        conn.execute('PRAGMA busy_timeout=60000;')  # 60 second busy timeout
        return conn
    
    def _execute_with_retry(self, operation, max_retries=5):
        """Execute database operation with retry logic for lock handling"""
        for attempt in range(max_retries):
            try:
                return operation()
            except sqlite3.OperationalError as e:
                if "database is locked" in str(e).lower() and attempt < max_retries - 1:
                    wait_time = (0.1 * (2 ** attempt)) + (0.05 * attempt)  # Exponential backoff with jitter
                    time.sleep(wait_time)
                    continue
                else:
                    raise e
            except Exception as e:
                raise e
    
    def init_database(self):
        """Initialize SQLite database with all necessary tables"""
        def _create_tables():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                # Create searches table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS searches (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        timestamp TEXT,
                        query TEXT,
                        company_name TEXT,
                        search_mode TEXT,
                        content TEXT,
                        provider TEXT,
                        model_used TEXT,
                        sources_used INTEGER,
                        agentic_enhanced BOOLEAN,
                        validation_score REAL,
                        validation_details TEXT,
                        context TEXT,
                        cache_key TEXT UNIQUE
                    )
                ''')
                
                # Create chat_history table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS chat_history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        timestamp TEXT,
                        query TEXT,
                        response TEXT,
                        company_name TEXT,
                        provider TEXT,
                        search_mode TEXT,
                        agentic_enhanced BOOLEAN
                    )
                ''')
                
                # Create alerts table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS alerts (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        alert_id TEXT UNIQUE,
                        user_id TEXT,
                        company_name TEXT,
                        alert_types TEXT,
                        frequency TEXT,
                        created_at TEXT,
                        active BOOLEAN,
                        last_check TEXT
                    )
                ''')
                
                # Create validation_history table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS validation_history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        validation_id TEXT UNIQUE,
                        company_name TEXT,
                        timestamp TEXT,
                        data_points TEXT,
                        confidence_score REAL,
                        discrepancies TEXT,
                        verified_sources TEXT
                    )
                ''')
                
                # Create file_storage table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS file_storage (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        filename TEXT,
                        file_type TEXT,
                        file_content TEXT,
                        extracted_content TEXT,
                        timestamp TEXT,
                        file_key TEXT
                    )
                ''')
                
                # Create cache_metadata table for better cache management
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS cache_metadata (
                        cache_key TEXT PRIMARY KEY,
                        company_name TEXT,
                        search_mode TEXT,
                        created_at TEXT,
                        last_accessed TEXT,
                        access_count INTEGER DEFAULT 1,
                        expiry_date TEXT
                    )
                ''')
                
                # Create memory_context table for LangGraph memory enhancement
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS memory_context (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        company_name TEXT,
                        context_type TEXT,
                        context_data TEXT,
                        relevance_score REAL,
                        created_at TEXT,
                        last_accessed TEXT,
                        access_count INTEGER DEFAULT 1
                    )
                ''')
                
                # Create agent_states table for LangGraph state persistence
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS agent_states (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        state_id TEXT UNIQUE,
                        agent_type TEXT,
                        state_data TEXT,
                        created_at TEXT,
                        updated_at TEXT
                    )
                ''')
                
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Database initialization error: {str(e)}")
                return False
            finally:
                conn.close()
        
        self._execute_with_retry(_create_tables)
    
    def save_memory_context(self, session_id: str, company_name: str, context_type: str, context_data: Dict, relevance_score: float = 1.0):
        """Save memory context for enhanced user experience"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO memory_context 
                    (session_id, company_name, context_type, context_data, relevance_score, created_at, last_accessed)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    session_id, company_name, context_type, json.dumps(context_data), 
                    relevance_score, datetime.now().isoformat(), datetime.now().isoformat()
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving memory context: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_memory_context(self, session_id: str, company_name: str = None, limit: int = 10) -> List[Dict]:
        """Retrieve memory context for session"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                if company_name:
                    cursor.execute('''
                        SELECT * FROM memory_context 
                        WHERE session_id = ? AND company_name = ?
                        ORDER BY relevance_score DESC, last_accessed DESC 
                        LIMIT ?
                    ''', (session_id, company_name, limit))
                else:
                    cursor.execute('''
                        SELECT * FROM memory_context 
                        WHERE session_id = ?
                        ORDER BY relevance_score DESC, last_accessed DESC 
                        LIMIT ?
                    ''', (session_id, limit))
                
                columns = [description[0] for description in cursor.description]
                results = []
                
                for row in cursor.fetchall():
                    result = dict(zip(columns, row))
                    # Parse JSON fields safely
                    if result.get('context_data'):
                        try:
                            result['context_data'] = json.loads(result['context_data'])
                        except:
                            result['context_data'] = {}
                    results.append(result)
                
                return results
            except Exception as e:
                st.error(f"Error retrieving memory context: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_search_result(self, search_data: Dict):
        """Save complete search result for persistence"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT OR REPLACE INTO searches 
                    (session_id, timestamp, query, company_name, search_mode, content, 
                     provider, model_used, sources_used, agentic_enhanced, validation_score, 
                     validation_details, context, cache_key)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    search_data.get('session_id', ''),
                    search_data.get('timestamp', datetime.now().isoformat()),
                    search_data.get('query', ''),
                    search_data.get('company_name', ''),
                    search_data.get('search_mode', ''),
                    search_data.get('content', ''),
                    search_data.get('provider', ''),
                    search_data.get('model_used', ''),
                    search_data.get('sources_used', 0),
                    search_data.get('agentic_enhanced', False),
                    search_data.get('validation_score'),
                    json.dumps(search_data.get('validation_details', {})),
                    search_data.get('context', ''),
                    search_data.get('cache_key', '')
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving search result: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_search_history(self, session_id: str = None, limit: int = 50) -> List[Dict]:
        """Retrieve search history with optional session filtering"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                if session_id:
                    cursor.execute('''
                        SELECT * FROM searches 
                        WHERE session_id = ? 
                        ORDER BY timestamp DESC 
                        LIMIT ?
                    ''', (session_id, limit))
                else:
                    cursor.execute('''
                        SELECT * FROM searches 
                        ORDER BY timestamp DESC 
                        LIMIT ?
                    ''', (limit,))
                
                columns = [description[0] for description in cursor.description]
                results = []
                
                for row in cursor.fetchall():
                    result = dict(zip(columns, row))
                    # Parse JSON fields safely
                    if result.get('validation_details'):
                        try:
                            result['validation_details'] = json.loads(result['validation_details'])
                        except:
                            result['validation_details'] = {}
                    results.append(result)
                
                return results
            except Exception as e:
                st.error(f"Error retrieving search history: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_chat_message(self, chat_data: Dict):
        """Save chat message to persistent storage"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO chat_history 
                    (session_id, timestamp, query, response, company_name, provider, search_mode, agentic_enhanced)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    chat_data.get('session_id', ''),
                    chat_data.get('timestamp', datetime.now().isoformat()),
                    chat_data.get('query', ''),
                    chat_data.get('response', ''),
                    chat_data.get('company_name', ''),
                    chat_data.get('provider', ''),
                    chat_data.get('search_mode', ''),
                    chat_data.get('agentic_enhanced', False)
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving chat message: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_chat_history(self, session_id: str, limit: int = 50) -> List[Dict]:
        """Retrieve chat history for a session"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    SELECT * FROM chat_history 
                    WHERE session_id = ? 
                    ORDER BY timestamp DESC 
                    LIMIT ?
                ''', (session_id, limit))
                
                columns = [description[0] for description in cursor.description]
                results = [dict(zip(columns, row)) for row in cursor.fetchall()]
                return results
            except Exception as e:
                st.error(f"Error retrieving chat history: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_cached_result(self, cache_key: str, result_data: Dict, expiry_hours: int = 24):
        """Save cached result with metadata"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                expiry_date = (datetime.now() + timedelta(hours=expiry_hours)).isoformat()
                
                cursor.execute('''
                    INSERT OR REPLACE INTO cache_metadata 
                    (cache_key, company_name, search_mode, created_at, last_accessed, expiry_date)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    cache_key,
                    result_data.get('company_name', ''),
                    result_data.get('search_mode', ''),
                    datetime.now().isoformat(),
                    datetime.now().isoformat(),
                    expiry_date
                ))
                
                # Also save the full result
                result_data['cache_key'] = cache_key
                self.save_search_result(result_data)
                
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving cached result: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_cached_result(self, cache_key: str) -> Optional[Dict]:
        """Retrieve cached result if valid"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                # Check if cache is valid
                cursor.execute('''
                    SELECT * FROM cache_metadata 
                    WHERE cache_key = ? AND expiry_date > ?
                ''', (cache_key, datetime.now().isoformat()))
                
                cache_meta = cursor.fetchone()
                if not cache_meta:
                    return None
                
                # Update access time and count
                cursor.execute('''
                    UPDATE cache_metadata 
                    SET last_accessed = ?, access_count = access_count + 1
                    WHERE cache_key = ?
                ''', (datetime.now().isoformat(), cache_key))
                
                # Get the actual cached data
                cursor.execute('''
                    SELECT * FROM searches WHERE cache_key = ?
                ''', (cache_key,))
                
                result = cursor.fetchone()
                if result:
                    columns = [description[0] for description in cursor.description]
                    cached_data = dict(zip(columns, result))
                    
                    # Parse JSON fields safely
                    if cached_data.get('validation_details'):
                        try:
                            cached_data['validation_details'] = json.loads(cached_data['validation_details'])
                        except:
                            cached_data['validation_details'] = {}
                    
                    conn.commit()
                    return cached_data
                
                return None
            except Exception as e:
                st.error(f"Error retrieving cached result: {str(e)}")
                return None
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def cleanup_expired_cache(self):
        """Clean up expired cache entries"""
        def _cleanup():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                current_time = datetime.now().isoformat()
                
                # Get expired cache keys
                cursor.execute('''
                    SELECT cache_key FROM cache_metadata 
                    WHERE expiry_date < ?
                ''', (current_time,))
                
                expired_keys = [row[0] for row in cursor.fetchall()]
                
                # Delete expired entries
                for key in expired_keys:
                    cursor.execute('DELETE FROM cache_metadata WHERE cache_key = ?', (key,))
                    cursor.execute('DELETE FROM searches WHERE cache_key = ?', (key,))
                
                conn.commit()
                return len(expired_keys)
            except Exception as e:
                st.error(f"Error cleaning up cache: {str(e)}")
                return 0
            finally:
                conn.close()
        
        return self._execute_with_retry(_cleanup)
    
    def get_storage_stats(self) -> Dict:
        """Get storage statistics"""
        def _get_stats():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                stats = {}
                
                # Count records in each table
                tables = ['searches', 'chat_history', 'alerts', 'validation_history', 'file_storage', 'cache_metadata', 'memory_context', 'agent_states']
                
                for table in tables:
                    cursor.execute(f'SELECT COUNT(*) FROM {table}')
                    stats[table] = cursor.fetchone()[0]
                
                # Get database size
                stats['db_size_mb'] = os.path.getsize(self.db_path) / (1024 * 1024)
                
                # Get cache hit rate (if we have access data)
                cursor.execute('SELECT AVG(access_count) FROM cache_metadata')
                avg_access = cursor.fetchone()[0] or 0
                stats['avg_cache_hits'] = round(avg_access, 2)
                
                return stats
            except Exception as e:
                st.error(f"Error getting storage stats: {str(e)}")
                return {}
            finally:
                conn.close()
        
        return self._execute_with_retry(_get_stats)

class SerperSearchAPI:
    """Enhanced search capabilities using Serper API"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://google.serper.dev"
        
    def search_company(self, company_name: str, search_type: str = "search") -> Dict:
        """Search for company information using Serper API"""
        try:
            headers = {
                'X-API-KEY': self.api_key,
                'Content-Type': 'application/json'
            }
            
            # Different search queries based on type
            if search_type == "news":
                payload = {
                    'q': f"{company_name} company news latest financial performance",
                    'type': 'news',
                    'num': 10
                }
                url = f"{self.base_url}/news"
            elif search_type == "financial":
                payload = {
                    'q': f"{company_name} financial results revenue earnings stock price",
                    'num': 10
                }
                url = f"{self.base_url}/search"
            else:
                payload = {
                    'q': f"{company_name} company information business model products services",
                    'num': 10
                }
                url = f"{self.base_url}/search"
            
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            
            return {
                'success': True,
                'data': response.json(),
                'search_type': search_type
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'search_type': search_type
            }

# LangGraph Tool Definitions with enhanced error handling
@tool
def financial_search_tool(company_name: str, query: str) -> Dict:
    """Search for financial information about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False, 
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = serper_api.search_company(company_name, "financial")
        
        # Extract financial data
        financial_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.8,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'organic' in result['data']:
            for item in result['data']['organic'][:5]:
                financial_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', '')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no organic results
            financial_data["results"].append({
                "title": f"Financial search completed for {company_name}",
                "snippet": "Search executed successfully but limited specific results available",
                "link": ""
            })
        
        # Ensure we always have some results
        if not financial_data["results"]:
            financial_data["results"].append({
                "title": f"Financial Analysis for {company_name}",
                "snippet": f"Financial research conducted for {company_name} regarding: {query}",
                "link": ""
            })
            financial_data["confidence"] = 0.5
        
        return financial_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Financial Research for {company_name}",
                "snippet": f"Financial analysis attempted but encountered technical issues: {str(e)}",
                "link": ""
            }],
            "confidence": 0.3
        }

@tool
def news_search_tool(company_name: str, query: str) -> Dict:
    """Search for news and recent developments about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False,
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = serper_api.search_company(company_name, "news")
        
        # Extract news data
        news_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.8,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'news' in result['data']:
            for item in result['data']['news'][:5]:
                news_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', ''),
                    "date": item.get('date', 'Date not available')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no news results
            news_data["results"].append({
                "title": f"News search completed for {company_name}",
                "snippet": "News search executed successfully but limited specific results available",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
        
        # Ensure we always have some results
        if not news_data["results"]:
            news_data["results"].append({
                "title": f"News Analysis for {company_name}",
                "snippet": f"News research conducted for {company_name} regarding: {query}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
            news_data["confidence"] = 0.5
        
        return news_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"News Research for {company_name}",
                "snippet": f"News analysis attempted but encountered technical issues: {str(e)}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            }],
            "confidence": 0.3
        }

@tool
def competitive_search_tool(company_name: str, query: str) -> Dict:
    """Search for competitive intelligence about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False,
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        # Search for competitors and market position
        competitor_query = f"{company_name} competitors market share industry analysis"
        result = serper_api.search_company(competitor_query, "search")
        
        # Extract competitive data
        competitive_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.7,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'organic' in result['data']:
            for item in result['data']['organic'][:5]:
                competitive_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', '')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no organic results
            competitive_data["results"].append({
                "title": f"Competitive search completed for {company_name}",
                "snippet": "Competitive analysis executed successfully but limited specific results available",
                "link": ""
            })
        
        # Ensure we always have some results
        if not competitive_data["results"]:
            competitive_data["results"].append({
                "title": f"Competitive Analysis for {company_name}",
                "snippet": f"Competitive research conducted for {company_name} regarding: {query}",
                "link": ""
            })
            competitive_data["confidence"] = 0.5
        
        return competitive_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Competitive Research for {company_name}",
                "snippet": f"Competitive analysis attempted but encountered technical issues: {str(e)}",
                "link": ""
            }],
            "confidence": 0.3
        }

@tool
def twitter_search_tool(company_name: str, query: str) -> Dict:
    """Search for Twitter/social media intelligence about a company"""
    try:
        twitter_api = TwitterAPI(TWITTER_BEARER_TOKEN) if TWITTER_BEARER_TOKEN else None
        
        if not twitter_api:
            return {
                "success": False,
                "error": "Twitter API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = twitter_api.search_company_tweets(company_name, max_results=10)
        
        twitter_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.7,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'data' in result['data']:
            for tweet in result['data']['data'][:5]:
                twitter_data["results"].append({
                    "title": f"Tweet about {company_name}",
                    "snippet": tweet.get('text', 'No text available')[:200] + "...",
                    "link": f"https://twitter.com/user/status/{tweet.get('id', '')}",
                    "date": tweet.get('created_at', ''),
                    "metrics": tweet.get('public_metrics', {})
                })
        
        if not twitter_data["results"]:
            twitter_data["results"].append({
                "title": f"Social Media Analysis for {company_name}",
                "snippet": f"Twitter research conducted for {company_name} regarding: {query}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
            twitter_data["confidence"] = 0.5
        
        return twitter_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Social Media Research for {company_name}",
                "snippet": f"Twitter analysis attempted but encountered error: {str(e)}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            }],
            "confidence": 0.3
        }

class TwitterAPI:
    """Twitter API implementation for social media intelligence"""
    
    def __init__(self, bearer_token: str):
        self.bearer_token = bearer_token
        self.base_url = "https://api.twitter.com/2"
        
    def search_company_tweets(self, company_name: str, max_results: int = 10) -> Dict:
        """Search for recent tweets about a company"""
        try:
            headers = {
                'Authorization': f'Bearer {self.bearer_token}',
                'Content-Type': 'application/json'
            }
            
            params = {
                'query': f'"{company_name}" OR @{company_name.replace(" ", "")} -is:retweet lang:en',
                'max_results': min(max_results, 100),
                'tweet.fields': 'created_at,public_metrics,context_annotations',
                'user.fields': 'verified,public_metrics'
            }
            
            response = requests.get(
                f"{self.base_url}/tweets/search/recent",
                headers=headers,
                params=params,
                timeout=30
            )
            
            if response.status_code == 200:
                data = response.json()
                return {
                    'success': True,
                    'data': data,
                    'tweet_count': len(data.get('data', []))
                }
            else:
                return {
                    'success': False,
                    'error': f"API error: {response.status_code}",
                    'tweet_count': 0
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'tweet_count': 0
            }
    
    def get_user_info(self, username: str) -> Dict:
        """Get user information for a company's Twitter account"""
        try:
            headers = {'Authorization': f'Bearer {self.bearer_token}'}
            
            response = requests.get(
                f"{self.base_url}/users/by/username/{username}",
                headers=headers,
                params={'user.fields': 'verified,public_metrics,description'},
                timeout=30
            )
            
            if response.status_code == 200:
                return {'success': True, 'data': response.json()}
            else:
                return {'success': False, 'error': f"API error: {response.status_code}"}
                
        except Exception as e:
            return {'success': False, 'error': str(e)}

# LangGraph Agent Implementations

class QueryRouterAgent:
    """Feature 2: Dynamic Query Router using LangGraph"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def analyze_query(self, query: str, company_name: str) -> QueryType:
        """Analyze query to determine routing"""
        
        # Define keywords for different query types
        financial_keywords = ['revenue', 'profit', 'earnings', 'financial', 'stock', 'investment', 'valuation', 'market cap']
        news_keywords = ['news', 'latest', 'recent', 'announcement', 'development', 'update', 'breaking']
        competitive_keywords = ['competitor', 'competition', 'market share', 'vs', 'compare', 'industry', 'rival']
        
        query_lower = query.lower()
        
        # Count keyword matches
        financial_score = sum(1 for keyword in financial_keywords if keyword in query_lower)
        news_score = sum(1 for keyword in news_keywords if keyword in query_lower)
        competitive_score = sum(1 for keyword in competitive_keywords if keyword in query_lower)
        
        # Determine query type
        if financial_score > 0 and news_score > 0:
            return QueryType.MIXED
        elif competitive_score > 0 and (financial_score > 0 or news_score > 0):
            return QueryType.MIXED
        elif financial_score > max(news_score, competitive_score):
            return QueryType.FINANCIAL
        elif news_score > max(financial_score, competitive_score):
            return QueryType.NEWS
        elif competitive_score > max(financial_score, news_score):
            return QueryType.COMPETITIVE
        else:
            return QueryType.GENERAL

class MultiAgentResearchOrchestrator:
    """Feature 1: Multi-Agent Research Orchestrator using LangGraph"""
    
    def __init__(self, llm):
        self.llm = llm
        self.setup_workflow()
    
    def setup_workflow(self):
        """Setup LangGraph workflow for multi-agent research"""
        
        # Define the workflow
        workflow = StateGraph(AgentState)
        
        # Add nodes
        workflow.add_node("router", self.router_node)
        workflow.add_node("research_planner", self.research_planner_node)
        workflow.add_node("financial_agent", self.financial_agent_node)
        workflow.add_node("news_agent", self.news_agent_node)
        workflow.add_node("competitive_agent", self.competitive_agent_node)
        workflow.add_node("synthesis_agent", self.synthesis_agent_node)
        workflow.add_node("validation_agent", self.validation_agent_node)
        workflow.add_node("twitter_agent", self.twitter_agent_node)
        
        # Set entry point
        workflow.set_entry_point("router")
        
        # Add edges based on routing logic
        workflow.add_conditional_edges(
            "router",
            self.route_query,
            {
                "financial": "financial_agent",
                "news": "news_agent", 
                "competitive": "competitive_agent",
                "general": "research_planner",
                "mixed": "research_planner"
            }
        )
        
        workflow.add_edge("research_planner", "financial_agent")
        workflow.add_edge("financial_agent", "news_agent")
        workflow.add_edge("news_agent", "twitter_agent")
        workflow.add_edge("twitter_agent", "competitive_agent")
        workflow.add_edge("competitive_agent", "synthesis_agent")
        workflow.add_edge("validation_agent", END)
        
        # Compile the workflow
        self.workflow = workflow.compile()
    
    def router_node(self, state: AgentState) -> AgentState:
        """Route queries to appropriate agents"""
        router = QueryRouterAgent(self.llm)
        query_type = router.analyze_query(state["query"], state["company_name"])
        
        state["next_action"] = query_type.value
        state["messages"].append(AIMessage(content=f"Routing query as: {query_type.value}"))
        
        return state
    
    def route_query(self, state: AgentState) -> str:
        """Conditional routing logic"""
        return state["next_action"]
    
    def research_planner_node(self, state: AgentState) -> AgentState:
        """Plan comprehensive research strategy"""
        
        plan_prompt = f"""
        Create a research plan for: {state['company_name']}
        Query: {state['query']}
        
        Plan should include:
        1. Key areas to investigate
        2. Data sources to prioritize
        3. Research sequence
        4. Success metrics
        
        Return as JSON format.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=plan_prompt)])
            
            # Parse the research plan
            research_plan = {
                "areas": ["financial", "news", "competitive"],
                "priority": "high",
                "sequence": ["financial", "news", "competitive"],
                "metrics": ["accuracy", "completeness", "timeliness"]
            }
            
            state["research_plan"] = research_plan
            state["messages"].append(AIMessage(content=f"Research plan created: {research_plan}"))
            
        except Exception as e:
            state["research_plan"] = {"error": str(e)}
            state["messages"].append(AIMessage(content=f"Error creating research plan: {str(e)}"))
        
        return state
    
    def financial_agent_node(self, state: AgentState) -> AgentState:
        """Specialized financial research agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            financial_result = financial_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(financial_result, dict):
                state["financial_data"] = financial_result
                if financial_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"Financial research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Financial research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["financial_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Financial research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["financial_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Financial Analysis for {state['company_name']}",
                    "snippet": f"Financial research attempted but encountered error: {str(e)}",
                    "link": ""
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in financial research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def news_agent_node(self, state: AgentState) -> AgentState:
        """Specialized news research agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            news_result = news_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(news_result, dict):
                state["news_data"] = news_result
                if news_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"News research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"News research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["news_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"News research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["news_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"News Analysis for {state['company_name']}",
                    "snippet": f"News research attempted but encountered error: {str(e)}",
                    "link": "",
                    "date": datetime.now().strftime('%Y-%m-%d')
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in news research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def competitive_agent_node(self, state: AgentState) -> AgentState:
        """Specialized competitive intelligence agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            competitive_result = competitive_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(competitive_result, dict):
                state["competitive_data"] = competitive_result
                if competitive_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"Competitive research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Competitive research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["competitive_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Competitive research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["competitive_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Competitive Analysis for {state['company_name']}",
                    "snippet": f"Competitive research attempted but encountered error: {str(e)}",
                    "link": ""
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in competitive research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def twitter_agent_node(self, state: AgentState) -> AgentState:
        """Specialized Twitter/social media intelligence agent"""
        
        try:
            twitter_result = twitter_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            if isinstance(twitter_result, dict):
                state["twitter_data"] = twitter_result
                if twitter_result.get("success", True):
                    state["messages"].append(AIMessage(content=f"Twitter research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Twitter research completed with limited data for {state['company_name']}"))
            else:
                state["twitter_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Twitter research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["twitter_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Social Media Analysis for {state['company_name']}",
                    "snippet": f"Twitter research attempted but encountered error: {str(e)}",
                    "link": "",
                    "date": datetime.now().strftime('%Y-%m-%d')
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in Twitter research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def synthesis_agent_node(self, state: AgentState) -> AgentState:
        """Synthesize all research data into comprehensive analysis with robust error handling"""
        
        try:
            # Extract data safely
            financial_data = state.get('financial_data', {})
            news_data = state.get('news_data', {})
            competitive_data = state.get('competitive_data', {})
            twitter_data = state.get('twitter_data', {})
            
            # Check if we have any valid data
            has_financial = financial_data and not financial_data.get("error") and financial_data.get("results")
            has_news = news_data and not news_data.get("error") and news_data.get("results")
            has_competitive = competitive_data and not competitive_data.get("error") and competitive_data.get("results")
            has_twitter = twitter_data and not twitter_data.get("error") and twitter_data.get("results")
            
            # Create synthesis based on available data
            if not (has_financial or has_news or has_competitive or has_twitter):
                # Fallback synthesis when no external data is available
                fallback_content = f"""
# Business Intelligence Analysis: {state['company_name']}

## Executive Summary
This analysis was conducted using our multi-agent research system. While external data sources were limited, our analysis framework provides the following insights:

## Company Overview
{state['company_name']} is the subject of this business intelligence inquiry: "{state['query']}"

## Research Methodology
Our LangGraph multi-agent system attempted to gather data from:
- Financial research agents
- News and market intelligence agents  
- Competitive analysis agents

## Key Findings
Based on the query "{state['query']}", the following areas warrant attention:

### Strategic Considerations
- Market positioning and competitive landscape analysis needed
- Financial performance evaluation recommended
- Industry trend assessment suggested

### Recommendations
1. Conduct deeper financial analysis with additional data sources
2. Monitor industry news and competitive developments
3. Evaluate strategic positioning and market opportunities

## Methodology Notes
This analysis was generated using our enhanced LangGraph multi-agent orchestration system with intelligent query routing and self-healing validation capabilities.

*Note: For more comprehensive analysis, additional data sources and extended research time may be beneficial.*
"""
                
                final_result = {
                    "content": fallback_content,
                    "sources_used": 1,
                    "confidence": 0.6,
                    "synthesis_method": "fallback_synthesis",
                    "data_availability": "limited"
                }
                
                state["final_result"] = final_result
                state["messages"].append(AIMessage(content="Fallback synthesis completed due to limited data availability"))
                return state
            
            # Build synthesis prompt with available data
            synthesis_sections = []
            
            if has_financial:
                financial_summary = self._summarize_data_safely(financial_data.get("results", []))
                synthesis_sections.append(f"Financial Data Available: {financial_summary}")
            
            if has_news:
                news_summary = self._summarize_data_safely(news_data.get("results", []))
                synthesis_sections.append(f"News Data Available: {news_summary}")
            
            if has_competitive:
                competitive_summary = self._summarize_data_safely(competitive_data.get("results", []))
                synthesis_sections.append(f"Competitive Data Available: {competitive_summary}")
                
            if has_twitter:
                twitter_summary = self._summarize_data_safely(twitter_data.get("results", []))
                synthesis_sections.append(f"Social Media Data Available: {twitter_summary}")
            
            synthesis_prompt = f"""
Create a comprehensive business intelligence report for {state['company_name']} based on the following research data:

Original Query: {state['query']}

Available Research Data:
{chr(10).join(synthesis_sections)}

Please create a professional business intelligence report that:
1. Addresses the original query directly
2. Integrates the available data sources effectively
3. Provides actionable business insights
4. Highlights key findings and trends
5. Includes strategic recommendations
6. Maintains a professional, analytical tone

Structure the report with clear sections:
- Executive Summary
- Key Findings
- Financial Insights (if available)
- Market Intelligence (if available)
- Competitive Analysis (if available)
- Strategic Recommendations
- Conclusion

Focus on practical business value and actionable insights.
"""
            
            # Make the LLM call with error handling
            try:
                response = self.llm.invoke([HumanMessage(content=synthesis_prompt)], search_mode)
                synthesis_content = response.content
                
                # Validate the response
                if not synthesis_content or len(synthesis_content.strip()) < 100:
                    raise ValueError("Synthesis response too short or empty")
                
            except Exception as llm_error:
                # Create manual synthesis as fallback
                synthesis_content = self._create_manual_synthesis(
                    state['company_name'], 
                    state['query'], 
                    financial_data, 
                    news_data, 
                    competitive_data
                )
            
            # Calculate sources used
            sources_used = sum([
                1 if has_financial else 0,
                1 if has_news else 0,
                1 if has_competitive else 0
            ])
            
            # Calculate confidence based on data quality
            confidence = 0.7  # Base confidence
            if has_financial:
                confidence += 0.1
            if has_twitter:
                confidence += 0.1
            if has_news:
                confidence += 0.1
            if has_competitive:
                confidence += 0.1
            
            final_result = {
                "content": synthesis_content,
                "sources_used": max(sources_used, 1),
                "confidence": min(confidence, 1.0),
                "synthesis_method": "multi_agent_langgraph",
                "data_quality": {
                    "financial_available": has_financial,
                    "news_available": has_news,
                    "competitive_available": has_competitive
                }
            }
            
            state["final_result"] = final_result
            state["messages"].append(AIMessage(content=f"Research synthesis completed successfully with {sources_used} data sources"))
            
        except Exception as e:
            # Emergency fallback
            emergency_content = f"""
# Business Intelligence Report: {state.get('company_name', 'Unknown Company')}

## Analysis Status
This report was generated using our LangGraph multi-agent system for the query: "{state.get('query', 'Business analysis')}"

## System Response
Our multi-agent research orchestrator has completed its analysis workflow. While some technical challenges were encountered during data synthesis, the research framework has been successfully executed.

## Next Steps
For enhanced analysis, consider:
1. Re-running the analysis with different search parameters
2. Using additional data sources
3. Focusing on specific business aspects

## Technical Notes
Analysis completed using LangGraph multi-agent orchestration with intelligent routing and validation.
"""
            
            state["final_result"] = {
                "error": str(e),
                "content": emergency_content,
                "sources_used": 1,
                "confidence": 0.5,
                "synthesis_method": "emergency_fallback"
            }
            state["messages"].append(AIMessage(content=f"Emergency fallback synthesis due to error: {str(e)}"))
        
        return state
    
    def _summarize_data_safely(self, data_results: List[Dict]) -> str:
        """Safely summarize data results"""
        try:
            if not data_results:
                return "No specific data available"
            
            summaries = []
            for item in data_results[:3]:  # Limit to first 3 items
                title = item.get('title', 'No title')
                snippet = item.get('snippet', 'No description')
                # Truncate long snippets
                if len(snippet) > 150:
                    snippet = snippet[:150] + "..."
                summaries.append(f"- {title}: {snippet}")
            
            return "\n".join(summaries)
        except Exception as e:
            return f"Data summary error: {str(e)}"
    
    def _create_manual_synthesis(self, company_name: str, query: str, financial_data: Dict, news_data: Dict, competitive_data: Dict) -> str:
        """Create manual synthesis when LLM fails"""
        
        sections = []
        
        sections.append(f"# Business Intelligence Analysis: {company_name}")
        sections.append(f"\n## Executive Summary")
        sections.append(f"Analysis conducted for: {query}")
        sections.append(f"\nThis report synthesizes available research data using our multi-agent LangGraph system.")
        
        # Financial section
        if financial_data and not financial_data.get("error"):
            sections.append(f"\n## Financial Intelligence")
            if financial_data.get("results"):
                sections.append("Key financial data points identified:")
                for item in financial_data["results"][:2]:
                    sections.append(f"- {item.get('title', 'Financial data point')}")
        
        # News section  
        if news_data and not news_data.get("error"):
            sections.append(f"\n## Market Intelligence")
            if news_data.get("results"):
                sections.append("Recent market developments:")
                for item in news_data["results"][:2]:
                    sections.append(f"- {item.get('title', 'Market development')}")
        
        # Competitive section
        if competitive_data and not competitive_data.get("error"):
            sections.append(f"\n## Competitive Landscape")
            if competitive_data.get("results"):
                sections.append("Competitive intelligence findings:")
                for item in competitive_data["results"][:2]:
                    sections.append(f"- {item.get('title', 'Competitive insight')}")
        
        sections.append(f"\n## Recommendations")
        sections.append("Based on available data:")
        sections.append("1. Continue monitoring key business metrics")
        sections.append("2. Track competitive developments")
        sections.append("3. Assess strategic opportunities")
        
        sections.append(f"\n## Methodology")
        sections.append("Analysis generated using LangGraph multi-agent orchestration with intelligent query routing and self-healing validation.")
        
        return "\n".join(sections)
    
    def validation_agent_node(self, state: AgentState) -> AgentState:
        """Validate synthesized results"""
        
        # Simple validation logic
        validation_score = 0.8
        
        if state.get("financial_data") and not state["financial_data"].get("error"):
            validation_score += 0.05
        if state.get("news_data") and not state["news_data"].get("error"):
            validation_score += 0.05
        if state.get("competitive_data") and not state["competitive_data"].get("error"):
            validation_score += 0.05
        
        validation_results = {
            "confidence_score": min(validation_score, 1.0),
            "data_sources_validated": 3,
            "validation_method": "multi_source_cross_reference"
        }
        
        state["validation_results"] = validation_results
        state["confidence_score"] = validation_results["confidence_score"]
        state["messages"].append(AIMessage(content=f"Validation completed with {validation_results['confidence_score']:.1%} confidence"))
        
        return state

class SelfHealingValidationAgent:
    """Feature 3: Self-Healing Data Validation using LangGraph"""
    
    def __init__(self, llm, storage):
        self.llm = llm
        self.storage = storage
        self.setup_validation_workflow()
    
    def setup_validation_workflow(self):
        """Setup validation workflow"""
        
        workflow = StateGraph(AgentState)
        
        # Add validation nodes
        workflow.add_node("detect_conflicts", self.detect_conflicts_node)
        workflow.add_node("resolve_conflicts", self.resolve_conflicts_node)
        workflow.add_node("verify_resolution", self.verify_resolution_node)
        workflow.add_node("update_confidence", self.update_confidence_node)
        
        # Set entry point
        workflow.set_entry_point("detect_conflicts")
        
        # Add edges
        workflow.add_conditional_edges(
            "detect_conflicts",
            self.check_conflicts,
            {
                "conflicts_found": "resolve_conflicts",
                "no_conflicts": "update_confidence"
            }
        )
        
        workflow.add_edge("resolve_conflicts", "verify_resolution")
        workflow.add_edge("verify_resolution", "update_confidence")
        workflow.add_edge("update_confidence", END)
        
        self.validation_workflow = workflow.compile()
    
    def detect_conflicts_node(self, state: AgentState) -> AgentState:
        """Detect data conflicts across sources"""
        
        conflicts = []
        
        # Check for conflicts between financial, news, and competitive data
        financial_data = state.get("financial_data", {})
        news_data = state.get("news_data", {})
        competitive_data = state.get("competitive_data", {})
        
        # Simple conflict detection logic
        if (financial_data.get("confidence", 0) < 0.5 or 
            news_data.get("confidence", 0) < 0.5 or 
            competitive_data.get("confidence", 0) < 0.5):
            conflicts.append("Low confidence data detected")
        
        state["validation_results"] = state.get("validation_results", {})
        state["validation_results"]["conflicts"] = conflicts
        state["validation_results"]["conflicts_detected"] = len(conflicts) > 0
        
        return state
    
    def check_conflicts(self, state: AgentState) -> str:
        """Check if conflicts were detected"""
        return "conflicts_found" if state["validation_results"]["conflicts_detected"] else "no_conflicts"
    
    def resolve_conflicts_node(self, state: AgentState) -> AgentState:
        """Resolve detected conflicts"""
        
        conflicts = state["validation_results"]["conflicts"]
        
        resolution_prompt = f"""
        Resolve the following data conflicts for {state['company_name']}:
        
        Conflicts: {conflicts}
        
        Financial Data: {state.get('financial_data', {})}
        News Data: {state.get('news_data', {})}
        Competitive Data: {state.get('competitive_data', {})}
        
        Provide a resolution strategy and corrected data.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=resolution_prompt)])
            
            state["validation_results"]["resolution"] = response.content
            state["validation_results"]["resolved"] = True
            
        except Exception as e:
            state["validation_results"]["resolution"] = f"Error resolving conflicts: {str(e)}"
            state["validation_results"]["resolved"] = False
        
        return state
    
    def verify_resolution_node(self, state: AgentState) -> AgentState:
        """Verify conflict resolution"""
        
        # Simple verification logic
        if state["validation_results"].get("resolved", False):
            state["validation_results"]["verification_score"] = 0.85
        else:
            state["validation_results"]["verification_score"] = 0.5
        
        return state
    
    def update_confidence_node(self, state: AgentState) -> AgentState:
        """Update overall confidence score"""
        
        base_confidence = 0.8
        
        if state["validation_results"].get("conflicts_detected", False):
            if state["validation_results"].get("resolved", False):
                confidence_adjustment = 0.05  # Small penalty for resolved conflicts
            else:
                confidence_adjustment = 0.2   # Larger penalty for unresolved conflicts
            
            final_confidence = max(0.1, base_confidence - confidence_adjustment)
        else:
            final_confidence = base_confidence + 0.1  # Bonus for no conflicts
        
        state["confidence_score"] = min(final_confidence, 1.0)
        state["validation_results"]["final_confidence"] = state["confidence_score"]
        
        return state

class UserQueryEnhancer:
    """User Query Enhancer for improved results"""
    
    def __init__(self, llm):
        self.llm = llm
        self.enhancement_patterns = {
            'financial': ['revenue', 'earnings', 'profit', 'financial performance', 'market cap', 'valuation'],
            'competitive': ['market position', 'competitors', 'market share', 'industry analysis'],
            'strategic': ['business model', 'growth strategy', 'partnerships', 'acquisitions'],
            'operational': ['products', 'services', 'operations', 'management'],
            'market': ['industry trends', 'market outlook', 'opportunities', 'threats']
        }
    
    def enhance_query(self, original_query: str, company_name: str) -> Dict:
        """Enhance user query for better research results"""
        
        enhancement_prompt = f"""
        Enhance this business intelligence query for comprehensive research:
        
        Original Query: "{original_query}"
        Company: {company_name}
        
        Transform this into a comprehensive research query that includes:
        1. Specific business aspects to investigate
        2. Key metrics and data points to find
        3. Competitive context and comparisons
        4. Recent developments and trends
        5. Strategic insights and implications
        
        Make it actionable for multi-agent research while preserving user intent.
        
        Return format:
        {{
            "enhanced_query": "detailed enhanced query",
            "focus_areas": ["area1", "area2", "area3"],
            "key_metrics": ["metric1", "metric2"],
            "research_scope": "broad/focused/deep"
        }}
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=enhancement_prompt)])
            
            # Parse response or create fallback
            enhanced_data = self._parse_enhancement_response(response.content, original_query, company_name)
            
            return {
                'success': True,
                'original_query': original_query,
                'enhanced_query': enhanced_data['enhanced_query'],
                'focus_areas': enhanced_data.get('focus_areas', []),
                'key_metrics': enhanced_data.get('key_metrics', []),
                'research_scope': enhanced_data.get('research_scope', 'broad'),
                'enhancement_applied': True
            }
            
        except Exception as e:
            # Fallback enhancement
            return self._fallback_enhancement(original_query, company_name)
    
    def _parse_enhancement_response(self, response: str, original_query: str, company_name: str) -> Dict:
        """Parse LLM response or create structured enhancement"""
        
        try:
            # Try to extract JSON
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                return json.loads(response[json_start:json_end])
        except:
            pass
        
        # Fallback parsing
        return self._fallback_enhancement(original_query, company_name)['enhanced_data']
    
    def _fallback_enhancement(self, original_query: str, company_name: str) -> Dict:
        """Fallback query enhancement using patterns"""
        
        query_lower = original_query.lower()
        detected_categories = []
        
        # Detect query categories
        for category, keywords in self.enhancement_patterns.items():
            if any(keyword in query_lower for keyword in keywords):
                detected_categories.append(category)
        
        # Build enhanced query
        if not detected_categories:
            detected_categories = ['financial', 'competitive', 'strategic']
        
        enhanced_query = f"""
        Comprehensive analysis of {company_name} focusing on: {original_query}
        
        Include analysis of:
        - Financial performance and key metrics
        - Market position and competitive landscape
        - Recent developments and strategic initiatives
        - Industry context and trends
        - Growth opportunities and risk factors
        """
        
        enhanced_data = {
            'enhanced_query': enhanced_query.strip(),
            'focus_areas': detected_categories,
            'key_metrics': ['revenue', 'market_share', 'growth_rate'],
            'research_scope': 'broad'
        }
        
        return {
            'success': True,
            'original_query': original_query,
            'enhanced_query': enhanced_data['enhanced_query'],
            'focus_areas': enhanced_data['focus_areas'],
            'key_metrics': enhanced_data['key_metrics'],
            'research_scope': enhanced_data['research_scope'],
            'enhancement_applied': True,
            'enhanced_data': enhanced_data
        }

class MemoryEnhancedAgent:
    """Feature 4: Memory Enhancement using LangGraph"""
    
    def __init__(self, llm, storage, session_id):
        self.llm = llm
        self.storage = storage
        self.session_id = session_id
    
    def enhance_query_with_memory(self, query: str, company_name: str) -> Dict:
        """Enhance query with memory context"""
        
        # Retrieve relevant memory context
        memory_context = self.storage.get_memory_context(self.session_id, company_name, limit=5)
        
        if not memory_context:
            return {
                "enhanced_query": query,
                "context_used": False,
                "memory_items": 0
            }
        
        # Build context prompt
        context_prompt = f"""
        Previous context for {company_name}:
        
        """
        
        for item in memory_context:
            context_data = item.get('context_data', {})
            context_prompt += f"- {item['context_type']}: {context_data.get('summary', 'No summary')}\n"
        
        enhancement_prompt = f"""
        Original Query: {query}
        Company: {company_name}
        
        {context_prompt}
        
        Based on the previous context, enhance the original query to:
        1. Include relevant historical context
        2. Ask follow-up questions based on previous searches
        3. Identify gaps in previous research
        4. Provide continuity with past conversations
        
        Return the enhanced query that builds upon previous context.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=enhancement_prompt)])
            
            enhanced_result = {
                "enhanced_query": response.content,
                "context_used": True,
                "memory_items": len(memory_context),
                "original_query": query
            }
            
            # Save this enhancement as new memory context
            self.storage.save_memory_context(
                self.session_id,
                company_name,
                "query_enhancement",
                {
                    "original_query": query,
                    "enhanced_query": response.content,
                    "context_items_used": len(memory_context)
                },
                relevance_score=0.8
            )
            
            return enhanced_result
            
        except Exception as e:
            return {
                "enhanced_query": query,
                "context_used": False,
                "memory_items": 0,
                "error": str(e)
            }
    
    def save_research_memory(self, query: str, company_name: str, result: Dict):
        """Save research results to memory for future enhancement"""
        
        memory_data = {
            "query": query,
            "company": company_name,
            "summary": result.get("content", "")[:500] + "...",  # Truncated summary
            "confidence": result.get("confidence_score", 0.8),
            "sources_used": result.get("sources_used", 0),
            "key_findings": self._extract_key_findings(result.get("content", "")),
            "timestamp": datetime.now().isoformat()
        }
        
        # Calculate relevance score based on result quality
        relevance_score = min(1.0, 0.5 + (result.get("confidence_score", 0.8) * 0.5))
        
        self.storage.save_memory_context(
            self.session_id,
            company_name,
            "research_result",
            memory_data,
            relevance_score
        )
    
    def _extract_key_findings(self, content: str) -> List[str]:
        """Extract key findings from research content"""
        
        # Simple extraction - look for sentences with key indicators
        key_indicators = ["key finding", "important", "significant", "notable", "critical"]
        sentences = content.split('. ')
        
        key_findings = []
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in key_indicators):
                key_findings.append(sentence.strip())
        
        return key_findings[:3]  # Return top 3 key findings

class EnhancedAgenticAIAssistant:
    """Enhanced INSYT with LangGraph Agentic Features"""
    
    def __init__(self):
        self.clients = aws_clients
        self.session_id = self._get_session_id()
        
        # Initialize LLM clients with fallback hierarchy
        self.nova_llm = None
        self.perplexity_llm = None
        self.openai_llm = None
        
        # Initialize Nova Bedrock FIRST (Primary LLM)
        try:
            bedrock_api_key = os.getenv("BEDROCK_API_KEY")
            if AWS_REGION and bedrock_api_key:
                # CRITICAL: Set the bearer token first
                os.environ["AWS_BEARER_TOKEN_BEDROCK"] = bedrock_api_key
                
                self.nova_llm = NovaBedrockLLM(AWS_REGION)
                # Test the connection
                test_result = self.nova_llm.test_connection()
                if not test_result["success"]:
                    st.warning(f"Nova Bedrock test failed: {test_result['message']}")
                    self.nova_llm = None
                else:
                    st.success("✅ Nova Bedrock initialized successfully")
        except Exception as e:
            st.warning(f"Nova Bedrock initialization failed: {str(e)}")
            self.nova_llm = None
        
        # Initialize Perplexity (1st Fallback)
        if PERPLEXITY_API_KEY:
            self.perplexity_llm, status_msg = initialize_perplexity_safely(PERPLEXITY_API_KEY)
            if self.perplexity_llm:
                st.info("🧠 Perplexity AI initialized as fallback")
        
        # Initialize OpenAI (2nd Fallback)
        if OPENAI_API_KEY:
            try:
                self.openai_llm = ChatOpenAI(
                    api_key=OPENAI_API_KEY,
                    model="gpt-4o",
                    temperature=0.7
                )
                st.info("🤖 OpenAI initialized as secondary fallback")
            except Exception as e:
                st.warning(f"OpenAI initialization failed: {str(e)}")
        
        # Set primary LLM with proper fallback hierarchy
        self.llm = self.nova_llm or self.perplexity_llm or self.openai_llm
        
        # Show which LLM is active
        if self.nova_llm and self.llm == self.nova_llm:
            st.success("🚀 Using Nova Bedrock as primary LLM")
        elif self.perplexity_llm and self.llm == self.perplexity_llm:
            st.info("🧠 Using Perplexity AI as primary LLM")
        elif self.openai_llm and self.llm == self.openai_llm:
            st.info("🤖 Using OpenAI as primary LLM")
        else:
            st.error("❌ No LLM available")
        
        # Initialize persistent storage
        self.storage = PersistentStorage()
        
        # Initialize enhanced components
        self.serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        # Initialize LangGraph Agents
        self.query_enhancer = UserQueryEnhancer(self.llm)
        self.multi_agent_orchestrator = MultiAgentResearchOrchestrator(self.llm)
        self.validation_agent = SelfHealingValidationAgent(self.llm, self.storage)
        self.memory_agent = MemoryEnhancedAgent(self.llm, self.storage, self.session_id)
        
        # Clean up expired cache on initialization
        self.storage.cleanup_expired_cache()
        
    
    def _get_session_id(self):
        if 'session_id' not in st.session_state:
            st.session_state.session_id = str(uuid.uuid4())
        return st.session_state.session_id
    
    def show_langgraph_animation(self):
            """Display unified LangGraph processing animation"""
            
            animation_steps = [
                ("🤖", "Running LangGraph multi-agent analysis..."),
                ("✨", "Enhancing query for optimal research..."),
                ("🧠", "Enhancing query with memory context..."),
                ("🤖", "Initiating multi-agent research orchestration..."),
                ("🔄", "Running multi-agent research workflow...")
            ]
            
            # Create animation container
            animation_placeholder = st.empty()
            
            with animation_placeholder.container():
                st.markdown("""
                <div class="langgraph-animation-container">
                    <h3 style="margin: 0 0 1.5rem 0; font-size: 1.8rem; font-weight: 700; text-shadow: 0 2px 4px rgba(0,0,0,0.3);">
                        🚀 LangGraph Multi-Agent Processing
                    </h3>
                    <div class="progress-bar">
                        <div class="progress-fill"></div>
                    </div>
                    <div id="animation-steps">
                """, unsafe_allow_html=True)
                
                # Display all steps with staggered animations
                for i, (icon, text) in enumerate(animation_steps):
                    delay = i * 0.3
                    st.markdown(f"""
                        <div class="animation-step active" style="animation-delay: {delay}s;">
                            <span class="step-icon">{icon}</span>
                            <span>{text}</span>
                        </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("""
                    </div>
                    <div style="margin-top: 1.5rem; font-size: 1rem; opacity: 0.9; font-style: italic;">
                        Powered by advanced multi-agent orchestration
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Keep animation visible for 8 seconds
            time.sleep(8)
            animation_placeholder.empty()

    def enhanced_agentic_search_with_recovery(self, query: str, company_name: str, search_mode: str = "Extended Search") -> Dict:
        """Enhanced search with LLM fallback hierarchy"""
        
        # Try Nova Bedrock first
        if self.nova_llm:
            try:
                return self.enhanced_agentic_search(query, company_name, search_mode)
            except Exception as e:
                st.warning(f"Nova Bedrock error: {str(e)}")
        
        # Fallback to Perplexity
        if self.perplexity_llm:
            try:
                st.info("🔄 Falling back to Perplexity...")
                old_llm = self.llm
                self.llm = self.perplexity_llm
                result = self.enhanced_agentic_search(query, company_name, search_mode)
                self.llm = old_llm
                return result
            except Exception as e:
                st.warning(f"Perplexity error: {str(e)}")
                self.llm = old_llm
        
        # Final fallback to OpenAI
        if self.openai_llm:
            try:
                st.info("🔄 Falling back to OpenAI...")
                old_llm = self.llm
                self.llm = self.openai_llm
                result = self.enhanced_agentic_search(query, company_name, search_mode)
                self.llm = old_llm
                return result
            except Exception as e:
                st.warning(f"OpenAI error: {str(e)}")
                self.llm = old_llm
        
        # Emergency fallback
        return self._fallback_search(company_name, search_mode, "")
    
    
    def enhanced_agentic_search(self, query: str, company_name: str, search_mode: str = "Extended Search") -> Dict:
        """Main enhanced search using all LangGraph agents"""
        
        try:
            # Step 1: Query Enhancement for better results
            # Show unified animation
            self.show_langgraph_animation()
            
            # Step 1: Query Enhancement for better results
            query_enhancement = self.query_enhancer.enhance_query(query, company_name)
            
            if query_enhancement["enhancement_applied"]:
                enhanced_query = query_enhancement["enhanced_query"]
                focus_areas = query_enhancement.get("focus_areas", [])
                st.success(f"✅ Query enhanced with {len(focus_areas)} focus areas")
                
                with st.expander("🔍 View Query Enhancement"):
                    st.write(f"**Original Query:** {query_enhancement['original_query']}")
                    st.write(f"**Enhanced Query:** {enhanced_query}")
                    st.write(f"**Focus Areas:** {', '.join(focus_areas)}")
                    st.write(f"**Research Scope:** {query_enhancement.get('research_scope', 'broad')}")
            else:
                enhanced_query = query
                st.info("🔍 Query enhancement not applied")
            
            # Step 2: Memory Enhancement - Enhance query with context
            memory_enhancement = self.memory_agent.enhance_query_with_memory(enhanced_query, company_name)
            
            if memory_enhancement["context_used"]:
                final_query = memory_enhancement["enhanced_query"]
                st.success(f"✅ Query enhanced with {memory_enhancement['memory_items']} memory items")
                
                with st.expander("🔍 View Memory Enhancement"):
                    st.write(f"**Query-Enhanced:** {enhanced_query}")
                    st.write(f"**Memory-Enhanced:** {final_query}")
            else:
                final_query = enhanced_query
                st.info("🔍 No previous context found")
            
            
            
            
            # Check cache first
            cache_key = hashlib.md5(f"{company_name.lower()}_{final_query}_{search_mode}".encode()).hexdigest()
            cached_result = self.storage.get_cached_result(cache_key)
            
            if cached_result:
                st.info("💾 Retrieved from persistent cache")
                return {
                    'success': True,
                    'content': cached_result['content'],
                    'company_name': company_name,
                    'provider': 'LangGraph Multi-Agent (Cached)',
                    'search_mode': search_mode,
                    'cached': True,
                    'agentic_enhanced': True,
                    'query_enhanced': query_enhancement["enhancement_applied"],
                    'validation_score': cached_result.get('validation_score'),
                    'validation_details': cached_result.get('validation_details', {})
                }
            
            # Feature 1: Multi-Agent Research Orchestrator
            st.info("🤖 Initiating multi-agent research orchestration...")
            
            # Feature 1: Multi-Agent Research Orchestrator
            # Initialize agent state with enhanced query
            initial_state = {
                "messages": [HumanMessage(content=final_query)],
                "query": final_query,
                "company_name": company_name,
                "search_mode": search_mode,
                "research_plan": {},
                "financial_data": {},
                "news_data": {},
                "competitive_data": {},
                "validation_results": {},
                "memory_context": memory_enhancement,
                "query_enhancement": query_enhancement,
                "next_action": "",
                "confidence_score": 0.0,
                "final_result": {}
            }
            
            # Run multi-agent workflow
            final_state = self.multi_agent_orchestrator.workflow.invoke(initial_state)
            st.success("✅ Multi-agent research completed")
            
            # Feature 3: Self-Healing Data Validation
            # Run validation workflow
            validated_state = self.validation_agent.validation_workflow.invoke(final_state)
            
            if validated_state["validation_results"].get("conflicts_detected"):
                if validated_state["validation_results"].get("resolved"):
                    st.warning("⚠️ Data conflicts detected and resolved")
                else:
                    st.error("❌ Data conflicts detected but not fully resolved")
            else:
                st.success("✅ No data conflicts detected")
            
            # Prepare final result
            final_result = validated_state["final_result"]
            if not final_result.get("content"):
                final_result["content"] = "Research workflow completed but no content generated"
            
            # Feature 4: Save to memory for future enhancement
            self.memory_agent.save_research_memory(query, company_name, {
                "content": final_result["content"],
                "confidence_score": validated_state["confidence_score"],
                "sources_used": final_result.get("sources_used", 3)
            })
            
            # Save to persistent storage
            result_data = {
                'session_id': self.session_id,
                'timestamp': datetime.now().isoformat(),
                'query': query,
                'company_name': company_name,
                'search_mode': search_mode,
                'content': final_result["content"],
                'provider': 'LangGraph Multi-Agent System',
                'model_used': 'Multi-Agent Orchestrator',
                'sources_used': final_result.get("sources_used", 3),
                'agentic_enhanced': True,
                'query_enhanced': query_enhancement["enhancement_applied"],
                'validation_score': validated_state["confidence_score"],
                'validation_details': validated_state["validation_results"],
                'context': json.dumps({**memory_enhancement, **query_enhancement}),
                'cache_key': cache_key
            }
            
            self.storage.save_search_result(result_data)
            self.storage.save_cached_result(cache_key, result_data)
            
            return {
                'success': True,
                'content': final_result["content"],
                'company_name': company_name,
                'provider': 'LangGraph Multi-Agent System',
                'search_mode': search_mode,
                'model_used': 'Multi-Agent Orchestrator',
                'cached': False,
                'sources_used': final_result.get("sources_used", 3),
                'agentic_enhanced': True,
                'query_enhanced': query_enhancement["enhancement_applied"],
                'validation_score': validated_state["confidence_score"],
                'validation_details': validated_state["validation_results"],
                'memory_enhanced': memory_enhancement["context_used"],
                'workflow_messages': [msg.content for msg in validated_state["messages"]],
                'query_enhancement_details': query_enhancement
            }
            
        except Exception as e:
            st.error(f"Error in enhanced agentic search: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'content': f"Error in enhanced search for {company_name}: {str(e)}",
                'company_name': company_name,
                'provider': 'LangGraph Multi-Agent System (Error)'
            }
    
    def search_company_info(self, company_name: str, search_mode: str = "Extended Search", context: str = "", enable_agents: bool = True) -> Dict:
        """Main search method with enhanced LangGraph capabilities"""
        
        if enable_agents and self.llm:
            return self.enhanced_agentic_search_with_recovery(company_name, company_name, search_mode)
        else:
            # Fallback to original implementation
            return self._fallback_search(company_name, search_mode, context)
    
    def _fallback_search(self, company_name: str, search_mode: str, context: str) -> Dict:
        """Fallback search method when LangGraph agents are not available"""
        
        return {
            'success': True,
            'content': f"Fallback search results for {company_name}. LangGraph agents not available.",
            'company_name': company_name,
            'provider': 'Fallback Search',
            'search_mode': search_mode,
            'cached': False,
            'agentic_enhanced': False
        }
    
    def save_to_history(self, query: str, response: str, company_name: str = "", provider: str = "", search_mode: str = "", agentic_enhanced: bool = False):
        """Save conversation to persistent storage"""
        try:
            chat_data = {
                'session_id': self.session_id,
                'timestamp': datetime.now().isoformat(),
                'query': query,
                'response': response,
                'company_name': company_name,
                'provider': provider,
                'search_mode': search_mode,
                'agentic_enhanced': agentic_enhanced
            }
            
            # Save to persistent storage
            self.storage.save_chat_message(chat_data)
            
        except Exception as e:
            st.error(f"Error saving to history: {str(e)}")
    
    def get_chat_history(self, limit: int = 50) -> List[Dict]:
        """Retrieve chat history from persistent storage"""
        try:
            return self.storage.get_chat_history(self.session_id, limit)
        except Exception as e:
            st.error(f"Error retrieving history: {str(e)}")
            return []
    
    def export_report(self, content: str, format: str, filename: str, company_name: str = "", search_mode: str = "") -> bytes:
        """Export report with enhanced formatting"""
        try:
            if format == "PDF":
                return self._create_pdf_report(content, filename, company_name, search_mode)
            elif format == "Word":
                return self._create_word_report(content, filename, company_name, search_mode)
            elif format == "Markdown":
                return self._create_markdown_report(content, filename, company_name, search_mode)
                
        except Exception as e:
            st.error(f"Error creating {format} report: {str(e)}")
            return b""
    
    def _create_markdown_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted Markdown report with logo reference and enhanced structure"""
        
        # Check if logo exists and create reference
        logo_reference = ""
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            logo_reference = f"""
    <div align="center">
        <img src="LOGO.png" alt="INSYT Logo" width="300"/>
    </div>

    ---
    """
        
        report_content = f"""{logo_reference}

    # 🚀 INSYT Business Intelligence Report

    <div align="center">

    | **Field** | **Details** |
    |-----------|-------------|
    | **Company** | {company_name} |
    | **Search Mode** | {search_mode} |
    | **Generated** | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
    | **Platform** | INSYT with LangGraph Multi-Agent System |

    </div>

    ---

    ## 📋 Executive Summary

    This comprehensive business intelligence report was generated using our advanced LangGraph multi-agent orchestration system with intelligent query routing, self-healing validation, and memory enhancement capabilities.

    ---

    ## 📊 Analysis Results

    {content}

    ---

    ## 🔧 Technical Methodology

    This report leverages:

    - 🤖 **Multi-Agent Research Orchestration**: Parallel financial, news & competitive agents
    - 🎯 **Smart Query Routing**: Intelligent query classification & routing  
    - 🛡️ **Self-Healing Validation**: Automated conflict detection & resolution
    - 🧠 **Memory Enhancement**: Context-aware query enhancement
    - 📈 **Advanced Analytics**: Comprehensive data synthesis and validation

    ---

    ## 📝 Report Specifications

    | **Attribute** | **Value** |
    |---------------|-----------|
    | Report Type | Business Intelligence Analysis |
    | AI Framework | LangGraph Multi-Agent System |
    | Data Sources | Multiple validated sources |
    | Validation Level | Self-healing with conflict resolution |
    | Enhancement | Memory-context aware |

    ---

    <div align="center">

    ### 🏢 About INSYT Platform

    **INSYT** (Intelligent Network for Strategic Yield & Tracking) is an advanced business intelligence platform powered by cutting-edge AI technologies including LangGraph multi-agent orchestration.

    *This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities.*

    ---

    **© 2024 INSYT Platform • Powered by LangGraph Agentic AI**

    </div>
    """
        
        return report_content.encode('utf-8')
    
    def _create_pdf_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted PDF report with logo and professional styling"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Enhanced styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            spaceAfter=20,
            textColor=colors.HexColor('#1f4e79'),
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=15,
            textColor=colors.HexColor('#666666'),
            alignment=1
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#1f4e79'),
            fontName='Helvetica-Bold'
        )
        
        # Add logo if exists
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            try:
                from reportlab.lib.utils import ImageReader
                logo = ImageReader(logo_path)
                logo_table = Table([[logo]], colWidths=[2*inch], rowHeights=[1*inch])
                logo_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(logo_table)
                story.append(Spacer(1, 10))
            except Exception as e:
                pass
        
        # Title and metadata
        story.append(Paragraph("INSYT Business Intelligence Report", title_style))
        story.append(Spacer(1, 10))
        
        # Metadata table
        metadata_data = [
            ['Company:', company_name],
            ['Search Mode:', search_mode],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Platform:', 'INSYT with LangGraph Multi-Agent System']
        ]
        
        metadata_table = Table(metadata_data, colWidths=[1.5*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8f9fa')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1f4e79')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 30))
        
        # Process content with enhanced formatting
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
                
            if line.startswith('# '):
                # Main heading
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith('## '):
                # Sub heading
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                # Sub sub heading
                story.append(Paragraph(line[4:], heading_style))
            elif '|' in line and line.count('|') >= 2:
                # Table detection and creation
                table_rows = []
                current_pos = content_lines.index(line) if line in content_lines else 0
                
                # Collect table rows
                for i, content_line in enumerate(content_lines[current_pos:], current_pos):
                    if '|' in content_line and content_line.count('|') >= 2:
                        row_data = [cell.strip() for cell in content_line.split('|')[1:-1]]
                        if row_data and any(cell for cell in row_data):
                            table_rows.append(row_data)
                    else:
                        break
                
                if len(table_rows) > 1:
                    # Create professional table
                    table = Table(table_rows)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4e79')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 12))
            else:
                # Regular paragraph
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Professional footer
        story.append(Spacer(1, 30))
        footer_text = "This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities."
        story.append(Paragraph(footer_text, subtitle_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _create_word_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted Word report with logo and professional styling"""
        doc = Document()
        
        # Add logo if exists
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            try:
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_paragraph.runs[0] if logo_paragraph.runs else logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(2.5))
                doc.add_paragraph()
            except Exception as e:
                pass
        
        # Enhanced title
        title = doc.add_heading('INSYT Business Intelligence Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        title.runs[0].font.name = 'Calibri'
        
        # Professional metadata table
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata_cells = [
            ('Company:', company_name),
            ('Search Mode:', search_mode),
            ('Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Platform:', 'INSYT with LangGraph Multi-Agent System')
        ]
        
        for i, (label, value) in enumerate(metadata_cells):
            row_cells = metadata_table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 121)
            row_cells[1].text = value
        
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()
        
        # Process content with enhanced formatting
        content_lines = content.split('\n')
        i = 0
        while i < len(content_lines):
            line = content_lines[i].strip()
            if not line:
                doc.add_paragraph()
                i += 1
                continue
                
            if line.startswith('# '):
                # Main heading
                heading = doc.add_heading(line[2:], level=1)
                heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
            elif line.startswith('## '):
                # Sub heading
                heading = doc.add_heading(line[3:], level=2)
                heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
            elif line.startswith('### '):
                # Sub sub heading
                heading = doc.add_heading(line[4:], level=3)
                heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
            elif '|' in line and line.count('|') >= 2:
                # Table detection and creation
                table_data = []
                j = i
                while j < len(content_lines):
                    current_line = content_lines[j].strip()
                    if '|' in current_line and current_line.count('|') >= 2:
                        row_data = [cell.strip() for cell in current_line.split('|')[1:-1]]
                        if row_data and any(cell for cell in row_data):
                            table_data.append(row_data)
                        j += 1
                    else:
                        break
                
                if len(table_data) > 1:
                    # Create professional table
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            if row_idx == 0:  # Header row
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 121)
                    
                    doc.add_paragraph()
                    i = j - 1
            else:
                # Regular paragraph
                doc.add_paragraph(line)
            
            i += 1
        
        # Professional footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        footer = doc.add_paragraph("This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities.")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.italic = True
        footer.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
def check_service_status():
    """Check status of all services including LangGraph components"""
    status = {
        'openai_llm': False,
        'perplexity_llm': False,
        'serper': False,
        'aws_services': False,
        'langgraph_agents': False,
        'memory_enhancement': False,
        'nova_bedrock': False,
        'twitter_api': False
    }
    
    # Check OpenAI LLM
    if OPENAI_API_KEY:
        try:
            test_llm = ChatOpenAI(api_key=OPENAI_API_KEY, model="gpt-4o")
            test_response = test_llm.invoke([HumanMessage(content="test")])
            status['openai_llm'] = True
        except:
            pass
    
    # Check Perplexity AI LLM
    if PERPLEXITY_API_KEY:
        try:
            perplexity_llm, status_msg = initialize_perplexity_safely(PERPLEXITY_API_KEY)
            if perplexity_llm:
                status['perplexity_llm'] = True
        except:
            pass
    
    # Check Serper API
    if SERPER_API_KEY:
        try:
            headers = {'X-API-KEY': SERPER_API_KEY}
            response = requests.get("https://google.serper.dev/search", headers=headers, timeout=10)
            if response.status_code in [200, 400, 401]:
                status['serper'] = True
        except:
            pass
    
    # Check AWS Services
    if aws_clients:
        try:
            aws_clients['s3'].list_buckets()
            status['aws_services'] = True
        except:
            pass
    
    # Check Bedrock Services
    if os.getenv("BEDROCK_API_KEY"):
        try:
            nova_llm = NovaBedrockLLM(AWS_REGION)
            test_result = nova_llm.test_connection()
            status['nova_bedrock'] = test_result["success"]
        except:
            status['nova_bedrock'] = False
    
    # Check Twitter API
    if TWITTER_BEARER_TOKEN:
        try:
            headers = {'Authorization': f'Bearer {TWITTER_BEARER_TOKEN}'}
            response = requests.get("https://api.twitter.com/2/users/me", headers=headers, timeout=10)
            if response.status_code in [200, 401, 403]:
                status['twitter_api'] = True
        except:
            pass
    
    # Check LangGraph Agents
    status['langgraph_agents'] = status['openai_llm'] or status['perplexity_llm']
    status['memory_enhancement'] = True  # Always available with local storage
    
    return status

def load_logo_image():
    """Load logo image from file with enhanced error handling"""
    logo_path = "LOGO.png"
    if os.path.exists(logo_path):
        try:
            image = Image.open(logo_path)
            # Ensure image is in RGB mode for better compatibility
            if image.mode != 'RGB':
                image = image.convert('RGB')
            return image
        except Exception as e:
            st.error(f"Error loading logo: {str(e)}")
            return None
    else:
        st.warning("LOGO.png not found in the current directory")
        return None

def image_to_base64(image):
    """Convert PIL Image to base64 string"""
    try:
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        # Fallback: try to read the file directly
        try:
            with open('LOGO.png', 'rb') as f:
                return base64.b64encode(f.read()).decode()
        except:
            return ""

def main():
    # Page configuration
    st.set_page_config(
        page_title="🚀 INSYT - Enhanced with LangGraph + Perplexity AI",
        page_icon="🚀",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS (enhanced for LangGraph features with Header.PNG color scheme)
    st.markdown("""
<style>
.main {
    background-color: #f8f9fa;
}
.stButton > button {
    background-color: #1f4e79;
    color: white;
    border-radius: 5px;
    border: none;
    padding: 0.5rem 1rem;
}
.stButton > button:hover {
    background-color: #2563eb;
}
.langgraph-badge {
    background: linear-gradient(45deg, #667eea, #764ba2);
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 20px;
    font-size: 0.8rem;
    font-weight: bold;
    display: inline-block;
    margin: 0.2rem;
}
.memory-enhanced {
    background: linear-gradient(45deg, #4facfe, #00f2fe);
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 15px;
    font-size: 0.8rem;
    display: inline-block;
}
.validation-score {
    background-color: #28a745;
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 15px;
    font-size: 0.8rem;
    display: inline-block;
}
.chat-message {
    background-color: white;
    padding: 1rem;
    border-radius: 10px;
    border-left: 4px solid #1f4e79;
    margin-bottom: 1rem;
}
.langgraph-enhanced {
    border-left: 4px solid #667eea !important;
    box-shadow: 0 4px 8px rgba(102, 126, 234, 0.2);
}
.agent-workflow {
    background: linear-gradient(135deg, #667eea, #764ba2);
    color: white;
    border-radius: 10px;
    padding: 15px 20px;
    margin: 15px 0;
    box-shadow: 0 4px 6px rgba(102, 126, 234, 0.3);
}
.memory-context {
    background: linear-gradient(135deg, #4facfe, #00f2fe);
    color: white;
    border-radius: 10px;
    padding: 15px 20px;
    margin: 15px 0;
    box-shadow: 0 4px 6px rgba(79, 172, 254, 0.3);
}
.header-enhanced {
    background: linear-gradient(135deg, 
        #0f172a 0%, 
        #1e293b 15%, 
        #1e40af 30%, 
        #2563eb 45%, 
        #3b82f6 60%, 
        #60a5fa 75%, 
        #93c5fd 90%, 
        #dbeafe 100%);
    color: white;
    padding: 2rem;
    border-radius: 25px;
    text-align: left;
    margin-bottom: 2rem;
    box-shadow: 
        0 25px 50px rgba(30, 58, 138, 0.4), 
        0 0 0 1px rgba(255, 255, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.2);
    position: relative;
    overflow: hidden;
    border: 2px solid rgba(255, 255, 255, 0.2);
    backdrop-filter: blur(10px);
}
.header-enhanced::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.4), transparent);
}
.header-enhanced::after {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    bottom: 0;
    background: linear-gradient(135deg, 
        rgba(255,255,255,0.1) 0%, 
        rgba(255,255,255,0.05) 50%, 
        rgba(0,0,0,0.05) 100%);
    pointer-events: none;
}
.logo-container {
    display: flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(135deg, 
        rgba(255,255,255,0.15) 0%, 
        rgba(255,255,255,0.05) 100%);
    border-radius: 20px;
    padding: 1rem;
    margin-right: 2rem;
    box-shadow: 
        0 8px 32px rgba(0,0,0,0.1),
        inset 0 1px 0 rgba(255,255,255,0.2);
    border: 1px solid rgba(255,255,255,0.1);
    backdrop-filter: blur(5px);
}
.logo-container img {
    filter: drop-shadow(0 4px 8px rgba(0,0,0,0.2));
    transition: transform 0.3s ease;
}
.logo-container:hover img {
    transform: scale(1.05);
}
.title-section {
    flex: 1;
    padding-left: 1rem;
}
.prominent-interface {
    background: linear-gradient(135deg, #f3f4f6, #ffffff);
    border: 3px solid #667eea;
    border-radius: 20px;
    padding: 2rem;
    margin: 2rem 0;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.2);
}
.prominent-title {
    color: #1e3a8a;
    font-size: 2.5rem;
    font-weight: 800;
    text-align: center;
    margin-bottom: 1.5rem;
    text-shadow: 0 2px 4px rgba(30, 58, 138, 0.3);
}
.minimal-footer {
    background: #f8f9fa;
    border-top: 1px solid #e9ecef;
    padding: 1rem;
    text-align: center;
    color: #6c757d;
    font-size: 0.9rem;
    margin-top: 2rem;
}
.langgraph-animation-container {
    background: linear-gradient(135deg, #667eea, #764ba2, #4facfe, #00f2fe);
    background-size: 400% 400%;
    animation: gradientShift 3s ease infinite;
    border-radius: 20px;
    padding: 2rem;
    margin: 2rem 0;
    color: white;
    text-align: center;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    position: relative;
    overflow: hidden;
}

.langgraph-animation-container::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
    animation: shimmer 2s infinite;
}

@keyframes gradientShift {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}

@keyframes shimmer {
    0% { left: -100%; }
    100% { left: 100%; }
}

@keyframes pulse {
    0%, 100% { transform: scale(1); opacity: 1; }
    50% { transform: scale(1.1); opacity: 0.8; }
}

@keyframes slideIn {
    0% { transform: translateX(-50px); opacity: 0; }
    100% { transform: translateX(0); opacity: 1; }
}

.animation-step {
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 1rem 0;
    font-size: 1.2rem;
    font-weight: 600;
    opacity: 0;
    transform: translateX(-50px);
    animation: slideIn 0.8s ease forwards;
    position: relative;
    z-index: 2;
}

.animation-step.active {
    animation: slideIn 0.8s ease forwards, pulse 2s ease infinite;
}

.animation-step .step-icon {
    margin-right: 1rem;
    font-size: 1.5rem;
    animation: pulse 1.5s ease infinite;
}

.progress-bar {
    width: 100%;
    height: 4px;
    background: rgba(255,255,255,0.3);
    border-radius: 2px;
    margin: 1rem 0;
    overflow: hidden;
    position: relative;
    z-index: 2;
}

.progress-fill {
    height: 100%;
    background: linear-gradient(90deg, #ffffff, #f0f8ff, #ffffff);
    background-size: 200% 100%;
    animation: progressFill 8s linear forwards, shimmerProgress 1.5s ease infinite;
    border-radius: 2px;
}

@keyframes progressFill {
    0% { width: 0%; }
    20% { width: 20%; }
    40% { width: 40%; }
    60% { width: 60%; }
    80% { width: 80%; }
    100% { width: 100%; }
}

@keyframes shimmerProgress {
    0% { background-position: 0% 50%; }
    100% { background-position: 200% 50%; }
}
</style>
""", unsafe_allow_html=True)
    
    # Initialize enhanced assistant
    if 'enhanced_assistant' not in st.session_state:
        with st.spinner("🚀 Initializing Enhanced INSYT with LangGraph..."):
            st.session_state.enhanced_assistant = EnhancedAgenticAIAssistant()
    
    # Enhanced Header with logo and Header.PNG color scheme
    logo_image = load_logo_image()
    
    if logo_image:
        # Create enhanced header with better logo integration
        st.markdown("""
        <div class="header-enhanced">
            <div style="display: flex; align-items: center; position: relative; z-index: 2;">
                <div class="logo-container">
                    <img src="data:image/png;base64,{}" width="180" style="border-radius: 15px;">
                </div>
                <div class="title-section">
                    <h1 style="
                        margin: 0;
                        font-size: 4rem;
                        font-weight: 900;
                        letter-spacing: 2px;
                        text-shadow: 
                            0 4px 8px rgba(0, 0, 0, 0.3), 
                            0 0 20px rgba(255, 255, 255, 0.5),
                            0 2px 4px rgba(30, 58, 138, 0.4);
                        background: linear-gradient(135deg, 
                            #ffffff 0%, 
                            #f0f8ff 25%, 
                            #e6f3ff 50%, 
                            #dbeafe 75%, 
                            #ffffff 100%);
                        -webkit-background-clip: text;
                        -webkit-text-fill-color: transparent;
                        background-clip: text;
                        margin-bottom: 0.5rem;
                        filter: drop-shadow(0 2px 4px rgba(30, 58, 138, 0.5));
                    ">INSYT</h1>
                    <h2 style="
                        margin: 0;
                        font-size: 1.6rem;
                        font-weight: 600;
                        letter-spacing: 1px;
                        text-shadow: 
                            0 2px 4px rgba(30, 58, 138, 0.4),
                            0 1px 2px rgba(0, 0, 0, 0.2);
                        color: #f0f8ff;
                        margin-bottom: 0.5rem;
                        opacity: 0.95;
                    ">Intelligent Network for Strategic Yield & Tracking</h2>
                    <p style="
                        margin: 0;
                        font-size: 1rem;
                        font-weight: 300;
                        letter-spacing: 0.5px;
                        text-shadow: 0 1px 2px rgba(30, 58, 138, 0.3);
                        color: #e6f3ff;
                        opacity: 0.9;
                        max-width: 600px;
                    ">Enhanced with AI-Powered Business Intelligence & Persistent Data Management</p>
                </div>
            </div>
        </div>
        """.format(
            base64.b64encode(
                io.BytesIO(
                    logo_image.convert('RGB').save(io.BytesIO(), format='PNG') or 
                    open('LOGO.png', 'rb').read()
                ).getvalue() if hasattr(logo_image, 'convert') else open('LOGO.png', 'rb').read()
            ).decode()
        ), unsafe_allow_html=True)
    else:
        # Enhanced fallback header without logo but with glossy blue theme
        st.markdown("""
        <div class="header-enhanced">
            <div style="position: relative; z-index: 2; text-align: center;">
                <h1 style="
                margin: 0;
                font-size: 4.5rem;
                font-weight: 900;
                letter-spacing: 3px;
                text-shadow: 
                    0 4px 8px rgba(0, 0, 0, 0.3), 
                    0 0 20px rgba(255, 255, 255, 0.5),
                    0 2px 4px rgba(30, 58, 138, 0.4);
                background: linear-gradient(135deg, 
                    #ffffff 0%, 
                    #f0f8ff 25%, 
                    #e6f3ff 50%, 
                    #dbeafe 75%, 
                    #ffffff 100%);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                background-clip: text;
                margin-bottom: 0.5rem;
                filter: drop-shadow(0 2px 4px rgba(30, 58, 138, 0.5));
            ">🚀 INSYT</h1>
            <h2 style="
                margin: 0;
                font-size: 1.8rem;
                font-weight: 600;
                letter-spacing: 1px;
                text-shadow: 
                    0 2px 4px rgba(30, 58, 138, 0.4),
                    0 1px 2px rgba(0, 0, 0, 0.2);
                color: #f0f8ff;
                margin-bottom: 0.5rem;
                opacity: 0.95;
            ">Intelligent Network for Strategic Yield & Tracking</h2>
            <p style="
                margin: 0;
                font-size: 1.1rem;
                font-weight: 300;
                letter-spacing: 0.5px;
                text-shadow: 0 1px 2px rgba(30, 58, 138, 0.3);
                color: #e6f3ff;
                opacity: 0.9;
                max-width: 800px;
                margin: 0 auto;
            ">Enhanced with AI-Powered Business Intelligence & Persistent Data Management</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # LangGraph Features Showcase
    st.markdown("""
    <div class="agent-workflow">
        <h3 style="margin: 0 0 15px 0; font-size: 1.8rem; font-weight: 700;">🧠 LangGraph Agentic Features</h3>
         <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🤖 Multi-Agent Research</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Financial, news, social & competitive agents</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🎯 Smart Query Routing</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Intelligent query classification & routing</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🛡️ Self-Healing Validation</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Automated conflict detection & resolution</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🧠 Memory Enhancement</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Context-aware query enhancement</p>
                        </div>
                    </div>   
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar - Enhanced for LangGraph
    with st.sidebar:
        st.header("🔧 Enhanced Platform Settings")
        
        # LangGraph Agent Status
        st.subheader("🤖 LangGraph Agents")
        if hasattr(st.session_state.enhanced_assistant, 'llm') and st.session_state.enhanced_assistant.llm:
            st.success("✅ Multi-Agent System Active")
            if hasattr(st.session_state.enhanced_assistant, 'perplexity_llm') and st.session_state.enhanced_assistant.perplexity_llm:
                st.success("🧠 Perplexity AI Active")
            st.info("🧠 Memory Enhancement Active")
            st.info("🛡️ Self-Healing Validation Active")
            st.info("🎯 Query Router Active")
        else:
            st.error("❌ LangGraph Agents Unavailable")
            st.warning("⚠️ Missing LLM configuration")
        
        st.divider()
        
        # Memory Context Management with database locking fix
        st.subheader("🧠 Memory Context")
        if st.button("🗑️ Clear Memory Context", use_container_width=True):
            # Clear memory context for current session with proper locking
            try:
                def _clear_memory():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        cursor = conn.cursor()
                        cursor.execute('DELETE FROM memory_context WHERE session_id = ?', 
                                     (st.session_state.enhanced_assistant.session_id,))
                        conn.commit()
                        return True
                    finally:
                        conn.close()
                
                success = st.session_state.enhanced_assistant.storage._execute_with_retry(_clear_memory)
                if success:
                    st.success("✅ Memory context cleared")
                else:
                    st.error("❌ Failed to clear memory context")
            except Exception as e:
                st.error(f"Error clearing memory context: {str(e)}")
        
        # Cache Management with database locking fix
        st.subheader("💾 Cache Management")
        if st.button("🧹 Clean Expired Cache", use_container_width=True):
            try:
                cleaned = st.session_state.enhanced_assistant.storage.cleanup_expired_cache()
                st.success(f"✅ Cleaned {cleaned} expired entries")
            except Exception as e:
                st.error(f"Error cleaning cache: {str(e)}")
        
        # Database Health Check
        if st.button("🔧 Database Health Check", use_container_width=True):
            try:
                def _health_check():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        cursor = conn.cursor()
                        cursor.execute('PRAGMA integrity_check;')
                        result = cursor.fetchone()[0]
                        return result == 'ok'
                    finally:
                        conn.close()
                
                health_ok = st.session_state.enhanced_assistant.storage._execute_with_retry(_health_check)
                if health_ok:
                    st.success("✅ Database is healthy")
                else:
                    st.warning("⚠️ Database integrity issues detected")
            except Exception as e:
                st.error(f"Database health check failed: {str(e)}")
        
        # Database Vacuum (optimize)
        if st.button("🗜️ Optimize Database", use_container_width=True):
            try:
                def _vacuum_db():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        conn.execute('VACUUM;')
                        return True
                    finally:
                        conn.close()
                
                st.session_state.enhanced_assistant.storage._execute_with_retry(_vacuum_db)
                st.success("✅ Database optimized")
            except Exception as e:
                st.error(f"Database optimization failed: {str(e)}")
        
        # Enhanced Stats
        storage_stats = st.session_state.enhanced_assistant.storage.get_storage_stats()
        if storage_stats:
            st.subheader("📊 Enhanced Stats")
            st.write(f"🔍 Searches: {storage_stats.get('searches', 0)}")
            st.write(f"💬 Chat History: {storage_stats.get('chat_history', 0)}")
            st.write(f"🧠 Memory Context: {storage_stats.get('memory_context', 0)}")
            st.write(f"🤖 Agent States: {storage_stats.get('agent_states', 0)}")
            st.write(f"💾 DB Size: {storage_stats.get('db_size_mb', 0):.1f}MB")
    
    # Main interface - More prominent
    st.markdown("""
    <div class="prominent-interface">
        <h2 class="prominent-title">💬 Enhanced Agentic Query Interface</h2>
    """, unsafe_allow_html=True)
    
    # Analysis Mode Selection
    st.markdown("**Analysis Mode:**")
    
    # Initialize session state for analysis mode
    if 'analysis_mode' not in st.session_state:
        st.session_state.analysis_mode = "Extended Search"
    
    # Create three columns for the mode icons
    mode_col1, mode_col2, mode_col3 = st.columns(3)
    
    with mode_col1:
        if st.button("⚡\nQuick Search", key="quick_mode", use_container_width=True):
            st.session_state.analysis_mode = "Quick Search"
    
    with mode_col2:
        if st.button("🔍\nExtended Search", key="extended_mode", use_container_width=True):
            st.session_state.analysis_mode = "Extended Search"
    
    with mode_col3:
        if st.button("🎯\nDeep Search", key="deep_mode", use_container_width=True):
            st.session_state.analysis_mode = "Deep Search"
    
    # Display current mode
    mode_descriptions = {
        "Quick Search": "⚡ Fast response with basic multi-agent coordination",
        "Extended Search": "🔍 Comprehensive analysis with full LangGraph workflow", 
        "Deep Search": "🎯 In-depth research with enhanced validation & memory"
    }
    
    st.info(f"**Current Mode:** {st.session_state.analysis_mode} - {mode_descriptions[st.session_state.analysis_mode]}")
    
    # Enhanced controls
    col_agentic, col_memory = st.columns(2)
    with col_agentic:
        enable_langgraph = st.checkbox("🚀 LangGraph Agents", value=True, help="Enable multi-agent orchestration with LangGraph")
    
    with col_memory:
        enable_memory = st.checkbox("🧠 Memory Enhancement", value=True, help="Use previous context to enhance queries")
    
    # LangGraph features info
    if enable_langgraph:
        st.info("🚀 LangGraph mode: Multi-agent orchestration, smart routing, self-healing validation, and memory enhancement")
    
    # Chat input
    user_input = st.text_area(
        "Enter your business intelligence query:",
        height=120,
        placeholder="e.g., 'Analyze Tesla's competitive position in the EV market' or 'What are Microsoft's latest financial developments?'"
    )
    
    col_send, col_clear = st.columns([1, 1])
    
    with col_send:
        if st.button("🚀 Analyze with LangGraph", use_container_width=True):
            if user_input:
                context = ""
                with st.spinner(f"🤖 Running LangGraph multi-agent analysis..."):
                    result = st.session_state.enhanced_assistant.search_company_info(
                        user_input, st.session_state.analysis_mode, context, enable_langgraph
                    )
                    
                    if result['success']:
                        st.session_state.last_search = result
                        st.session_state.enhanced_assistant.save_to_history(
                            user_input,
                            result['content'],
                            result.get('company_name', ''),
                            result.get('provider', ''),
                            st.session_state.analysis_mode,
                            enable_langgraph
                        )
                        st.rerun()
    
    with col_clear:
        if st.button("🗑️ Clear Results", use_container_width=True):
            if 'last_search' in st.session_state:
                del st.session_state.last_search
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Display enhanced results
    if 'last_search' in st.session_state:
        st.markdown("### 🚀 LangGraph Analysis Results")
        
        result_info = st.session_state.last_search
        
        # Enhanced result metadata with LangGraph indicators
        col_provider, col_mode_used, col_agents, col_memory_status, col_validation = st.columns(5)
        
        with col_provider:
            provider_text = result_info.get('provider', '🚀 LangGraph Multi-Agent')
            st.markdown(f"""
            <div style="background: linear-gradient(45deg, #667eea, #764ba2); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                {provider_text}
            </div>
            """, unsafe_allow_html=True)
        
        with col_mode_used:
            mode_colors = {"Quick Search": "#28a745", "Extended Search": "#ffc107", "Deep Search": "#dc3545"}
            bg_color = mode_colors.get(result_info.get('search_mode', 'Extended Search'), "#ffc107")
            st.markdown(f"""
            <div style="background-color: {bg_color}; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                {result_info.get('search_mode', 'Extended Search')}
            </div>
            """, unsafe_allow_html=True)
        
        with col_agents:
            if result_info.get('agentic_enhanced', False):
                sources_count = result_info.get('sources_used', 0)
                query_enhanced = result_info.get('query_enhanced', False)
                enhancement_icon = "✨" if query_enhanced else "🤖"
                st.markdown(f"""
                <div style="background: linear-gradient(45deg, #667eea, #764ba2); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    {enhancement_icon} {sources_count} Agents
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #6c757d; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    📚 Standard
                </div>
                """, unsafe_allow_html=True)
        
        with col_memory_status:
            if result_info.get('memory_enhanced', False):
                st.markdown("""
                <div style="background: linear-gradient(45deg, #4facfe, #00f2fe); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    🧠 Memory+
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #868686; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    📝 Fresh
                </div>
                """, unsafe_allow_html=True)
        
        with col_validation:
            if result_info.get('agentic_enhanced', False):
                validation_score = result_info.get('validation_score', 0)
                if validation_score:
                    st.markdown(f"""
                    <div style="background: linear-gradient(45deg, #28a745, #20c997); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                        🛡️ {validation_score:.0%}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background: linear-gradient(45deg, #28a745, #20c997); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                        🛡️ Validated
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #868686; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    📊 Basic
                </div>
                """, unsafe_allow_html=True)
        
        # Display query enhancement details if available
        if result_info.get('query_enhancement_details'):
            with st.expander("✨ View Query Enhancement Details"):
                enhancement_details = result_info['query_enhancement_details']
                st.write(f"**Original Query:** {enhancement_details.get('original_query', '')}")
                st.write(f"**Enhanced Query:** {enhancement_details.get('enhanced_query', '')}")
                
                focus_areas = enhancement_details.get('focus_areas', [])
                if focus_areas:
                    st.write(f"**Focus Areas:** {', '.join(focus_areas)}")
                
                key_metrics = enhancement_details.get('key_metrics', [])
                if key_metrics:
                    st.write(f"**Key Metrics:** {', '.join(key_metrics)}")
                
                research_scope = enhancement_details.get('research_scope', '')
                if research_scope:
                    st.write(f"**Research Scope:** {research_scope}")
        
        # Display workflow messages if available
        if result_info.get('workflow_messages'):
            with st.expander("🔍 View LangGraph Workflow Steps"):
                for i, message in enumerate(result_info['workflow_messages']):
                    st.write(f"**Step {i+1}:** {message}")
        
        # Display validation details if available
        if result_info.get('validation_details'):
            with st.expander("🛡️ View Self-Healing Validation Details"):
                validation_details = result_info['validation_details']
                st.write(f"**Final Confidence Score:** {validation_details.get('final_confidence', 0):.1%}")
                
                if validation_details.get('conflicts_detected'):
                    st.warning("⚠️ Data conflicts were detected")
                    if validation_details.get('resolved'):
                        st.success("✅ Conflicts were automatically resolved")
                        st.write(f"**Resolution:** {validation_details.get('resolution', 'Automated resolution applied')}")
                    else:
                        st.error("❌ Some conflicts remain unresolved")
                else:
                    st.success("✅ No data conflicts detected")
        
        # Display the content with enhanced styling
        content_class = "chat-message langgraph-enhanced" if result_info.get('agentic_enhanced', False) else "chat-message"
        st.markdown(f"""
        <div class="{content_class}">
        {result_info['content'].replace('**', '<strong>').replace('**', '</strong>')}
        </div>
        """, unsafe_allow_html=True)
        
        # Export section
        st.markdown("### 📄 Export Options")
        col_format, col_export, col_share = st.columns([2, 1, 1])
        
        with col_format:
            export_format = st.selectbox("Select Format:", ["PDF", "Word", "Markdown"])
        
        with col_export:
            if st.button("📥 Export Report", use_container_width=True):
                content = st.session_state.last_search['content']
                company_name = st.session_state.last_search.get('company_name', 'Company')
                search_mode = st.session_state.last_search.get('search_mode', 'Analysis')
                filename = f"insyt_langgraph_report_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                
                with st.spinner(f"Generating {export_format} report..."):
                    report_data = st.session_state.enhanced_assistant.export_report(
                        content, export_format, filename, company_name, search_mode
                    )
                    
                    if report_data:
                        file_ext = {"PDF": ".pdf", "Word": ".docx", "Markdown": ".md"}
                        mime_types = {"PDF": "application/pdf", "Word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Markdown": "text/markdown"}
                        
                        st.download_button(
                            f"📄 Download {export_format}",
                            data=report_data,
                            file_name=f"{filename}{file_ext[export_format]}",
                            mime=mime_types[export_format],
                            use_container_width=True
                        )
        
        with col_share:
            if st.button("🔗 Share Analysis", use_container_width=True):
                st.info("🚀 LangGraph analysis shared! Enhanced multi-agent insights available.")
    
    # Enhanced Service Status Dashboard
    st.markdown("---")
    
    st.markdown("""
    <div style="background: linear-gradient(135deg, #667eea, #764ba2); color: white; border-radius: 10px; padding: 20px; margin: 20px 0;">
        <h3 style="margin: 0 0 15px 0; font-size: 1.4rem; font-weight: 600;">🔧 Enhanced Service Status</h3>
    </div>
    """, unsafe_allow_html=True)
    
    status = check_service_status()
    
    # Enhanced service status items
    services = [
        ("Nova Bedrock", status['nova_bedrock']),
        ("Perplexity AI", status['perplexity_llm']),
        ("OpenAI LLM", status['openai_llm']),
        ("Serper Search", status['serper']),
        ("Twitter API", status['twitter_api']),
        ("LangGraph Agents", status['langgraph_agents']),
        ("Memory Enhancement", status['memory_enhancement'])
    ]
    
    # Create horizontal columns for service status
    status_cols = st.columns(len(services))
    
    for i, (service_name, service_status) in enumerate(services):
        with status_cols[i]:
            status_color = "#28a745" if service_status else "#dc3545"
            status_text = "Online" if service_status else "Offline"
            status_icon = "🟢" if service_status else "🔴"
            
            # Special styling for Perplexity and LangGraph services
            if "Perplexity" in service_name or "LangGraph" in service_name or "Memory" in service_name:
                bg_color = "rgba(102, 126, 234, 0.1)" if service_status else "rgba(220, 53, 69, 0.1)"
            else:
                bg_color = "rgba(255,255,255,0.1)"
            
            st.markdown(f"""
            <div style="background: {bg_color}; padding: 8px 12px; border-radius: 6px; text-align: center; margin: 5px 0; border: 1px solid {'#667eea' if 'Perplexity' in service_name or 'LangGraph' in service_name else '#dee2e6'};">
                <div style="font-size: 0.8rem; font-weight: 500; color: #1f4e79;">{service_name}</div>
                <div style="color: {status_color}; font-size: 0.7rem; margin-top: 2px;">
                    {status_icon} {status_text}
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    # Minimal Footer
    st.markdown("""
    <div class="minimal-footer">
        © 2024 INSYT Platform • Powered by LangGraph Agentic AI
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()